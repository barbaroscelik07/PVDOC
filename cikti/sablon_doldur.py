"""
Şablon doldurma yardımcıları (python-docx).

Şablonu açıp içeriğini doldurmak için güvenli, yeniden kullanılabilir işlemler:
- metin değiştirme (run'lar bölünmüş olsa bile paragraf düzeyinde),
- tablo satırını biçimini koruyarak klonlama (dinamik satır ekleme),
- hücreye biçimli metin yazma.

Bu yaklaşım şablonun font/stil/boşluk/kenarlık biçimini korur; sıfırdan
üretimde kaybedilen "birebir görünüm" böyle elde edilir.
"""

from __future__ import annotations

import copy
from docx.table import _Row, Table
from docx.text.paragraph import Paragraph


def paragraf_metni_degistir(paragraf: Paragraph, eski: str, yeni: str) -> bool:
    """
    Paragraf içindeki 'eski' metni 'yeni' ile değiştirir. Metin birden çok
    run'a bölünmüş olabileceği için tüm paragraf metnini birleştirip ilk run'a
    yazar, kalan run'ları temizler (paragraf biçimi korunur).
    Değişiklik olduysa True döner.
    """
    tam = "".join(r.text for r in paragraf.runs)
    if eski not in tam:
        return False
    yeni_tam = tam.replace(eski, yeni)
    if paragraf.runs:
        paragraf.runs[0].text = yeni_tam
        for r in paragraf.runs[1:]:
            r.text = ""
    return True


def belgede_degistir(doc, eslemeler: dict[str, str]) -> None:
    """
    Tüm paragraflarda ve tablo hücrelerinde verilen metin eşlemelerini uygular.
    eslemeler: {'aranacak': 'yazılacak', ...}
    """
    def _isle(paragraflar):
        for p in paragraflar:
            for eski, yeni in eslemeler.items():
                paragraf_metni_degistir(p, eski, yeni)

    _isle(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _isle(cell.paragraphs)
    # üst/alt bilgi
    for section in doc.sections:
        for hf in (section.header, section.footer):
            _isle(hf.paragraphs)
            for t in hf.tables:
                for row in t.rows:
                    for cell in row.cells:
                        _isle(cell.paragraphs)


def satir_klonla(tablo: Table, kaynak_index: int = -1) -> _Row:
    """
    Tablodaki bir satırı (varsayılan: sonuncu) biçimiyle birlikte klonlayıp
    tablonun sonuna ekler ve yeni satırı döndürür.
    Hücre biçimi/kenarlık/gölge korunur; metin sonra yazılır.
    """
    kaynak = tablo.rows[kaynak_index]
    yeni_tr = copy.deepcopy(kaynak._tr)
    tablo._tbl.append(yeni_tr)
    return tablo.rows[-1]


def hucre_yaz(cell, metin: str, bold: bool | None = None) -> None:
    """
    Hücreye metin yazar; ilk run'ın biçimini (font) korur, kalan run'ları VE
    fazla paragrafları siler (eski şablon içeriği kalmaz).
    bold None ise mevcut bold durumu korunur. Yeni run Times New Roman olur.
    """
    from docx.shared import Pt as _Pt, RGBColor as _RGB
    p = cell.paragraphs[0]
    for ekstra in cell.paragraphs[1:]:
        ekstra._p.getparent().remove(ekstra._p)
    if p.runs:
        p.runs[0].text = str(metin)
        if bold is not None:
            p.runs[0].bold = bold
        for r in p.runs[1:]:
            r.text = ""
        # font yine de Times New Roman'a sabitle (Calibri kaçaklarını önle)
        if not p.runs[0].font.name:
            p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.color.rgb = _RGB(0, 0, 0)  # her zaman siyah
    else:
        r = p.add_run(str(metin))
        r.font.name = "Times New Roman"
        r.font.size = _Pt(12)
        r.font.color.rgb = _RGB(0, 0, 0)
        if bold is not None:
            r.bold = bold


def satiri_bosalt(row: _Row) -> None:
    """
    Bir satırın tüm hücrelerini boşaltır (klon sonrası temiz başlangıç).
    Ayrıca dikey/yatay hücre birleştirmelerini (vMerge/gridSpan) KALDIRIR;
    aksi halde klonlanan satır üstteki hücrenin değerini görsel olarak devralır
    ve Operasyon No/Operasyon sütunları kayar.
    """
    from docx.oxml.ns import qn
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""
        tcPr = cell._tc.find(qn("w:tcPr"))
        if tcPr is not None:
            for etiket in ("w:vMerge", "w:gridSpan"):
                el = tcPr.find(qn(etiket))
                if el is not None:
                    tcPr.remove(el)
