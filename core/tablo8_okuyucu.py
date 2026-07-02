"""
Tablo 8 (Bitmiş Ürün Serbest Bırakma Spesifikasyonları) Word okuyucu.

Kullanıcı tüm PVR/PVP Word dosyasını yükler; bu modül "Tablo 8" başlıklı
2 sütunlu (TESTLER | SPESİFİKASYONLAR) tabloyu bulur ve içeriği yapısal
olarak çözer:
  - Görünüş, Ortalama Ağırlık, Dağılma → tek test
  - Ağırlık Tekdüzeliği → 2 alt satır (sapabilir / sapmamalıdır)
  - Etkin madde başlıkları → o etkene ait Teşhis/Miktar Tayini
  - Dissolüsyon
  - İlgili Bileşikler → etkin madde 'e Ait' grupları + impuriteler
  - Mikrobiyolojik Kontrol → 3 alt satır
"""

from __future__ import annotations

from docx import Document
from docx.table import Table

from core.models import (
    Test, Spesifikasyon, LimitTuru, TabloTipi, EtkinMadde, Impurite,
)


def _norm(s: str) -> str:
    tr = str.maketrans({"ı": "i", "İ": "i", "ş": "s", "Ş": "s", "ğ": "g",
                        "ç": "c", "ö": "o", "ü": "u", "I": "i"})
    return (s or "").translate(tr).lower().strip()


def tablo8_bul(yol: str) -> Table | None:
    """Word dosyasında 'Tablo 8 ... Bitmiş Ürün' tablosunu bulur."""
    d = Document(yol)
    body = list(d.element.body)
    bulundu = False
    for el in body:
        if el.tag.endswith("}p"):
            txt = _norm("".join(el.itertext()))
            if "tablo" in txt and "8" in txt and "bitmis urun" in txt:
                bulundu = True
        elif el.tag.endswith("}tbl") and bulundu:
            return Table(el, d)
    # Yedek: 2 sütunlu, ilk satırı TESTLER/SPESİFİKASYON olan ilk tablo
    for tbl in d.tables:
        if len(tbl.columns) == 2 and tbl.rows:
            h0 = _norm(tbl.rows[0].cells[0].text)
            h1 = _norm(tbl.rows[0].cells[1].text)
            if "test" in h0 and "spesifikasyon" in h1:
                return tbl
    return None


_USTSIMGE = {"0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
             "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
             "+": "⁺", "-": "⁻", "n": "ⁿ"}


def _hucre_metni(cell) -> str:
    """
    Hücre metnini okur; superscript (üst simge) run'ları Unicode üst simgeye
    çevirir. Böylece 10³/10² düzleşip '103/102' olmaz.
    """
    from docx.oxml.ns import qn
    parcalar = []
    for p in cell.paragraphs:
        for r in p.runs:
            metin = r.text
            rPr = r._element.find(qn('w:rPr'))
            ust = False
            if rPr is not None:
                va = rPr.find(qn('w:vertAlign'))
                if va is not None and va.get(qn('w:val')) == "superscript":
                    ust = True
            if ust:
                metin = "".join(_USTSIMGE.get(c, c) for c in metin)
            parcalar.append(metin)
        parcalar.append("\n")
    return "".join(parcalar).strip()


def _sayi(metin: str):
    import re
    m = re.search(r"(\d+(?:[.,]\d+)?)", metin or "")
    return float(m.group(1).replace(",", ".")) if m else None


def tablo8_coz(yol: str) -> dict:
    """
    Tablo 8'i çözüp şu yapıyı döndürür:
      {
        "testler": [Test, ...],          # bitmiş ürün testleri (etken adı gömülü)
        "etkin_maddeler": [EtkinMadde],  # impurite grupları ile
        "bulundu": bool,
      }
    """
    t = tablo8_bul(yol)
    if t is None:
        return {"bulundu": False, "testler": [], "etkin_maddeler": []}

    testler: list[Test] = []
    etkenler: dict[str, EtkinMadde] = {}
    aktif_etken_imp = None     # İlgili Bileşikler altında aktif etken
    ilgili_bolumde = False
    enantiomerik_bolumde = False
    mikro_test = None
    agirlik_test = None
    aktif_etken_adi = None     # "Etkin madde 1" başlığı altındayız
    _bekleyen_boyar = []       # Boyar Madde gibi alt-satır spec'i bekleyen testler

    def _etken(ad):
        if ad not in etkenler:
            etkenler[ad] = EtkinMadde(ad=ad)
        return etkenler[ad]

    rows = t.rows[1:]  # başlık satırını atla
    for ri_idx, row in enumerate(rows):
        sol = _hucre_metni(row.cells[0])
        sag = _hucre_metni(row.cells[1])
        n = _norm(sol)
        if not sol and not sag:
            continue
        # bir sonraki dolu satırın sol hücresi (etkin madde başlığı tespiti için)
        sonraki_sol = ""
        for ileri in rows[ri_idx + 1:]:
            s = _hucre_metni(ileri.cells[0])
            if s:
                sonraki_sol = s
                break

        # --- Bölüm başlıkları ---
        if n == "ilgili bilesikler" or n.startswith("ilgili bilesik"):
            ilgili_bolumde = True
            enantiomerik_bolumde = False
            # Grup başlığı ("'e Ait") yoksa impuriteler son tanınan etkene gider.
            if aktif_etken_adi:
                aktif_etken_imp = _etken(aktif_etken_adi)
            elif etkenler:
                aktif_etken_imp = list(etkenler.values())[-1]
            else:
                aktif_etken_imp = None
            continue

        # --- Enantiomerik İmpurite (İlgili Bileşikler ile aynı yapı) ---
        if "enantiomerik" in n:
            enantiomerik_bolumde = True
            ilgili_bolumde = False
            if aktif_etken_adi:
                aktif_etken_imp = _etken(aktif_etken_adi)
            elif etkenler:
                aktif_etken_imp = list(etkenler.values())[-1]
            else:
                aktif_etken_imp = None
            continue

        if enantiomerik_bolumde:
            if sol.startswith(("—", "-", "–")):
                if aktif_etken_imp is not None:
                    ad = sol.lstrip("—-– ").strip()
                    te = "t.e" in _norm(sag)
                    aktif_etken_imp.enantiomerik.append(Impurite(
                        ad=ad, limit_metni=sag, maksimum_deger=_sayi(sag), te=te))
                    continue
            else:
                enantiomerik_bolumde = False

        if ilgili_bolumde:
            # "Etkin madde 1'e Ait" → grup başlığı
            if "e ait" in n or "a ait" in n or "'e ait" in sol.lower() or "'a ait" in sol.lower():
                ad = sol.split("'")[0].split("’")[0].strip()
                if not ad:
                    ad = sol.replace("e Ait", "").replace("a Ait", "").strip()
                aktif_etken_imp = _etken(ad)
                continue
            # impurite satırı (— ile başlar)
            if sol.startswith(("—", "-", "–")):
                if aktif_etken_imp is not None:
                    ad = sol.lstrip("—-– ").strip()
                    te = "t.e" in _norm(sag)
                    aktif_etken_imp.impuriteler.append(Impurite(
                        ad=ad, limit_metni=sag, maksimum_deger=_sayi(sag), te=te))
                    continue
            else:
                # — ile başlamayan satır → İlgili Bileşikler bölümü bitti
                # (yeni etkin madde başlığı veya düz test). Aşağıda normal işlenecek.
                ilgili_bolumde = False

        # --- Mikrobiyolojik ---
        if "mikrobiyolojik" in n:
            ilgili_bolumde = False
            mikro_test = Test(
                ad="Mikrobiyolojik Kontrol", tablo_tipi=TabloTipi.MATRIS,
                mikrobiyolojik=True,
                spesifikasyon=Spesifikasyon(limit_turu=LimitTuru.METIN, sabit_sonuc="Uygun"),
                alt_satirlar=[])
            testler.append(mikro_test)
            continue
        if mikro_test is not None and sol.startswith(("—", "-", "–")) and not ilgili_bolumde:
            mikro_test.alt_satirlar.append((sol, sag))
            continue

        # --- Ağırlık Tekdüzeliği / Ağırlık Sapması (başlık + 2 alt) ---
        if ("agirlik tekduzeligi" in n or "agirlik sapmasi" in n) and not sag:
            agirlik_test = Test(
                ad=sol.strip(), tablo_tipi=TabloTipi.AGIRLIK_TEKDUZELIGI,
                spesifikasyon=Spesifikasyon(limit_turu=LimitTuru.ARALIK, birim="mg"))
            testler.append(agirlik_test)
            continue
        if agirlik_test is not None and sol.startswith(("—", "-", "–")):
            if not agirlik_test.aciklama_etiketi:
                agirlik_test.aciklama_etiketi = sol
                agirlik_test.aciklama_spek = sag
            else:
                agirlik_test.aciklama2_etiketi = sol
                agirlik_test.aciklama2_spek = sag
                # alt/üst limiti açıklamadan türet
                sayilar = _alt_ust_metinden(sag) or _alt_ust_metinden(agirlik_test.aciklama_spek)
            continue
        else:
            if agirlik_test is not None and not sol.startswith(("—", "-", "–")):
                agirlik_test = None  # ağırlık bölümü bitti

        # --- Etkin madde başlığı ---
        # Yapısal tespit: sağ hücre boş + bir sonraki satır "—" ile başlayıp
        # Teşhis/Miktar içeriyorsa, bu satır bir etkin madde başlığıdır (ismi ne olursa olsun).
        sonraki_n = _norm(sonraki_sol.lstrip("—-– "))
        sonraki_alt = sonraki_sol.startswith(("—", "-", "–"))
        etken_basligi = (
            not sag and "ilgili" not in n and "dissol" not in n
            and "mikrobiyolojik" not in n and "agirlik" not in n
            and (
                n.startswith("etkin madde")
                or (sonraki_alt and ("teshis" in sonraki_n or "miktar" in sonraki_n))
            )
        )
        if etken_basligi:
            aktif_etken_adi = sol.strip()
            _etken(aktif_etken_adi)  # etkin maddeler sözlüğüne ekle
            continue

        # --- Alt satır: Teşhis / Miktar / İçerik Tekdüzeliği (aktif etken altında) ---
        if sol.startswith(("—", "-", "–")) and aktif_etken_adi:
            alt_ad = sol.lstrip("—-– ").strip()
            alt_ad = _ad_temizle(alt_ad)
            tam_ad = f"{aktif_etken_adi} {alt_ad}"
            tip = TabloTipi.IKI_NUMUNE if "miktar" in _norm(alt_ad) else TabloTipi.TEK_SONUC
            testler.append(_test_yap(tam_ad, sag, tip))
            continue

        # --- Düz testler ---
        if sol and not sol.startswith(("—", "-", "–")):
            ad_temiz = _ad_temizle(sol)
            # Etken başlığı altındaki "Dissolüsyon" o etkene bağlanır
            if "dissol" in n and aktif_etken_adi:
                testler.append(_test_yap(f"{aktif_etken_adi} {ad_temiz}", sag, TabloTipi.BOS_NOKTA))
                continue
            aktif_etken_adi = None
            tip = _tip_tahmin(n)
            yeni_test = _test_yap(ad_temiz, sag, tip)
            testler.append(yeni_test)
            # Boyar Madde gibi başlık(boş)+alt satırlı testler: alt satırları
            # alt_satirlar listesinde biriktir (İlgili Bileşikler benzeri yapı)
            if not sag and "boyar madde" in n:
                yeni_test.alt_satirlar = []
                _bekleyen_boyar.append(yeni_test)
            continue
        # Boyar Madde alt satırı (-Titanyum dioksit ...)
        if sol.startswith(("—", "-", "–")) and _bekleyen_boyar and not aktif_etken_adi:
            hedef = _bekleyen_boyar[-1]
            alt_ad = sol.lstrip("—-– ").strip()
            hedef.alt_satirlar.append((alt_ad, sag))
            # İlk alt satırı spesifikasyon metni olarak da tut (geri uyumluluk)
            if not hedef.spesifikasyon.spesifikasyon_metni:
                hedef.spesifikasyon.spesifikasyon_metni = sag
                hedef.spesifikasyon.sabit_sonuc = sag
            continue

    return {"bulundu": True, "testler": testler,
            "etkin_maddeler": list(etkenler.values())}


def _ad_temizle(ad: str) -> str:
    """Test adından parantezli açıklama ve satır sonlarını temizler."""
    import re
    ad = ad.replace("\n", " ").strip()
    # '(Kütle varyasyonuna göre)' gibi parantezli ekleri at
    ad = re.sub(r"\s*\([^)]*\)\s*", " ", ad).strip()
    return ad


def _alt_ust_metinden(metin):
    import re
    s = re.sub(r"(\d),(\d)", r"\1.\2", metin or "")
    return [float(x) for x in re.findall(r"\d+(?:\.\d+)?", s)]


def _tip_tahmin(n: str) -> TabloTipi:
    if "ortalama agirlik" in n:
        return TabloTipi.BOS_NOKTA
    if "dagilma" in n or "dissol" in n:
        return TabloTipi.BOS_NOKTA
    if "miktar" in n:
        return TabloTipi.IKI_NUMUNE
    return TabloTipi.TEK_SONUC


def _test_yap(ad: str, spek_metni: str, tip: TabloTipi) -> Test:
    sp = Spesifikasyon(limit_turu=LimitTuru.ARALIK, spesifikasyon_metni=spek_metni)
    # alt/üst limiti metinden türet
    sayilar = _alt_ust_metinden(spek_metni)
    if len(sayilar) >= 2:
        sp.alt_limit = sayilar[-2]
        sp.ust_limit = sayilar[-1]
        sp.alt_metin = str(sayilar[-2])
        sp.ust_metin = str(sayilar[-1])
        sp.hedef_deger = sayilar[0]
    elif len(sayilar) == 1:
        if "maksimum" in _norm(spek_metni) or "max" in _norm(spek_metni):
            sp.ust_limit = sayilar[0]; sp.maksimum_deger = sayilar[0]
        elif "minimum" in _norm(spek_metni) or "min" in _norm(spek_metni):
            sp.alt_limit = sayilar[0]; sp.minimum_deger = sayilar[0]
    # metinsel testler
    if _norm(ad).endswith("gorunus") or "teshis" in _norm(ad):
        sp.limit_turu = LimitTuru.METIN
        sp.sabit_sonuc = spek_metni
    # Miktar Tayini birimini ayıkla (Tablo 9 toleransı için): "500 mg/f.t ± %5" → "mg/f.t"
    if "miktar" in _norm(ad):
        import re
        m = re.search(r"\d+(?:[.,]\d+)?\s*([a-zA-Zğüşıöçİ./]+)", spek_metni or "")
        if m:
            sp.birim = m.group(1).strip()
    return Test(ad=ad, tablo_tipi=tip, spesifikasyon=sp)
