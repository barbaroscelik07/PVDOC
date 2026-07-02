"""
DOCX üretim motoru — ŞABLON DOLDURMA yaklaşımı.

PVP/PVR şablonu açılır, placeholder'lar kullanıcı verisiyle değiştirilir ve
dinamik tablolar (formül, kapsanan ürünler, risk, proses param, ekipman,
spesifikasyon, numune planı, PVR sonuçları) şablonun kendi satır biçimi
korunarak doldurulur.

Bu yaklaşım şablonun fontunu, satır boşluklarını, kenarlıklarını, kapak
sayfasını, içindekiler bölümünü ve gömülü resimleri (IBC numune resmi vb.)
BİREBİR korur — sıfırdan üretimde kaybedilen "taslakla aynı görünüm" budur.

Tablolar indeks yerine ilk-hücre metnine göre bulunur (şablon değişse de sağlam).
"""

from __future__ import annotations

import sys
import os
from pathlib import Path

from docx import Document

from core.models import (
    ProjeVerisi, UrunFormu, Test, TabloTipi, LimitTuru,
    SERI_SAYISI, NOKTA_ADLARI,
)
from core import veri_uretici as vu
from cikti.sablon_doldur import (
    belgede_degistir, satir_klonla, hucre_yaz, satiri_bosalt,
)


# ----------------------------------------------------------------------------
# Şablon yolu (PyInstaller uyumlu)
# ----------------------------------------------------------------------------

def _sablon_yolu(ad: str) -> Path:
    taban = getattr(sys, "_MEIPASS", os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    return Path(taban) / "cikti" / "sablonlar" / ad


# ----------------------------------------------------------------------------
# Yardımcılar
# ----------------------------------------------------------------------------

def _urun(proje: ProjeVerisi) -> str:
    return proje.dokuman.urun_adi or "Xxx Film Kaplı Tablet"


def _tablo_bul(doc, ilk_hucre_metni: str, sutun_sayisi: int | None = None):
    """İlk hücresi verilen metinle başlayan ilk tabloyu döndürür."""
    hedef = ilk_hucre_metni.strip().lower()
    for t in doc.tables:
        if not t.rows:
            continue
        ilk = t.rows[0].cells[0].text.strip().lower()
        if ilk.startswith(hedef):
            if sutun_sayisi is None or len(t.columns) == sutun_sayisi:
                return t
    return None


def _tablo_basliga_gore(doc, tablo_no: int):
    """
    'Tablo N ...' başlık paragrafından HEMEN SONRA gelen tabloyu döndürür.
    Benzer yapılı tabloları (Tablo 6 vs 8.2) kesin ayırt eder — en sağlam yöntem.
    """
    from docx.table import Table as _T
    from docx.text.paragraph import Paragraph as _P
    from docx.oxml.ns import qn as _qn

    onceki_baslik_no = None
    for child in doc.element.body.iterchildren():
        if child.tag == _qn("w:p"):
            txt = _P(child, doc).text.strip()
            # "Tablo 6 ..." veya "Tablo.6 ..." veya "Tablo 6\t..."
            low = txt.lower()
            if low.startswith("tablo"):
                # ilk sayıyı çek
                import re
                m = re.search(r"tablo[\s\.]*(\d+)", low)
                onceki_baslik_no = int(m.group(1)) if m else None
        elif child.tag == _qn("w:tbl"):
            if onceki_baslik_no == tablo_no:
                return _T(child, doc)
            onceki_baslik_no = None  # tablo geçildi, başlık tüketildi
    return None


def _seri_nolar(proje: ProjeVerisi) -> list[str]:
    out = []
    for i in range(SERI_SAYISI):
        sno = proje.seriler[i].seri_no if i < len(proje.seriler) else ""
        out.append(sno or f"YYY-P{i+1:02d}")
    return out


def _tablo_genislik_duzelt(tablo):
    """
    Tablo yapısı (tcW/tblGrid) şablonda zaten doğru ve Word'de düzgün render
    edilir. LibreOffice'in çok satırlı tablolardaki render bug'ı DOCX'i bozmaz;
    bu yüzden burada yapısal değişiklik yapmıyoruz (Word çıktısı korunur).
    """
    return


def _tablo_render_duzelt(tablo):
    """
    LibreOffice PDF render'ında çok satırlı tabloların sütun çökmesini önler:
    fixed layout + her hücreye tblGrid'den gelen sabit genişlik uygular.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = tablo._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        return
    genislikler = [int(gc.get(qn('w:w'))) for gc in grid.findall(qn('w:gridCol'))
                   if gc.get(qn('w:w'))]
    if not genislikler:
        return
    toplam = sum(genislikler)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(toplam)); tblW.set(qn('w:type'), 'dxa')
    # her hücreye sabit genişlik
    for row in tablo.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(genislikler):
                break
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(genislikler[ci])); tcW.set(qn('w:type'), 'dxa')


def _veri_satirlarini_ayarla(tablo, baslik_satir_sayisi: int, gereken_veri_satiri: int):
    """
    Tablodaki veri satırı sayısını 'gereken'e eşitler:
    - fazla satırları siler,
    - eksikse son veri satırını klonlar.
    Başlık satırları (ilk N) korunur. Döndürür: veri satırlarının indeks listesi.
    """
    mevcut_veri = len(tablo.rows) - baslik_satir_sayisi
    # fazlaları sil (sondan)
    while mevcut_veri > gereken_veri_satiri and mevcut_veri > 0:
        tr = tablo.rows[-1]._tr
        tr.getparent().remove(tr)
        mevcut_veri -= 1
    # eksikleri klonla
    while mevcut_veri < gereken_veri_satiri:
        yeni = satir_klonla(tablo, kaynak_index=-1)
        satiri_bosalt(yeni)
        mevcut_veri += 1
    return list(range(baslik_satir_sayisi, baslik_satir_sayisi + gereken_veri_satiri))


def _tablo_sabit_layout(tablo):
    """
    Tabloya sabit (fixed) layout ve tablo genişliği uygular; LibreOffice'in
    çok satırlı tablolarda sütunları çökertmesini önler.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = tablo._tbl
    tblPr = tbl.tblPr
    # tblLayout = fixed
    mevcut = tblPr.find(qn('w:tblLayout'))
    if mevcut is None:
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
    else:
        mevcut.set(qn('w:type'), 'fixed')
    # tblW = grid toplamı (dxa)
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        toplam = sum(int(gc.get(qn('w:w'))) for gc in grid.findall(qn('w:gridCol'))
                     if gc.get(qn('w:w')))
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
        tblW.set(qn('w:w'), str(toplam))
        tblW.set(qn('w:type'), 'dxa')


# ============================================================================
# PVP tablolarını doldur
# ============================================================================

def _doldur_formul(doc, proje: ProjeVerisi) -> None:
    t = _tablo_basliga_gore(doc, 1)  # Tablo 1
    if t is None or not proje.hammaddeler:
        return
    # {adet} placeholder'ı seri boyutundaki adet değeriyle değiştir (başlık hücresi)
    adet = proje.seriler[0].seri_boyutu_adet if proje.seriler else ""
    if adet:
        for row in t.rows:
            for cell in row.cells:
                if "{adet}" in cell.text:
                    from cikti.sablon_doldur import paragraf_metni_degistir
                    for p in cell.paragraphs:
                        paragraf_metni_degistir(p, "{adet}", adet)

    _USTLER = {1: "¹", 2: "²", 3: "³", 4: "⁴", 5: "⁵", 6: "⁶", 7: "⁷", 8: "⁸", 9: "⁹"}

    def _vf(x, ond):
        """Türkçe ondalık (virgül) biçimlendirme."""
        if x is None:
            return ""
        return f"{x:.{ond}f}".replace(".", ",")

    # --- Yıldız gruplarını belirle ---
    # Uçucu sıvılar TEK bir yıldız grubunu paylaşır (aynı 'Uçucudur' notu).
    # Kaplama materyali AYRI bir yıldız grubu. Yıldız sayısı, grubun birim
    # formülde İLK göründüğü sıraya göre atanır (ilk grup *, ikinci grup **).
    ilk_ucucu_idx = None
    ilk_kaplama_idx = None
    for i, h in enumerate(proje.hammaddeler):
        if getattr(h, "ucucu_sivi", False) and ilk_ucucu_idx is None:
            ilk_ucucu_idx = i
        if (getattr(h, "kaplama_yildiz", 0) or "kaplama materyali" in (h.ad or "").lower()) \
           and ilk_kaplama_idx is None:
            ilk_kaplama_idx = i
    # Gruplara yıldız sayısı ata (önce gelen 1 yıldız, sonraki 2)
    gruplar = []
    if ilk_ucucu_idx is not None:
        gruplar.append(("ucucu", ilk_ucucu_idx))
    if ilk_kaplama_idx is not None:
        gruplar.append(("kaplama", ilk_kaplama_idx))
    gruplar.sort(key=lambda x: x[1])
    yildiz_say = {}  # grup adı -> yıldız sayısı
    for n, (grup, _idx) in enumerate(gruplar, 1):
        yildiz_say[grup] = n

    son = len(proje.hammaddeler) - 1
    var_km = False
    var_uy = False
    idxler = _veri_satirlarini_ayarla(t, 1, len(proje.hammaddeler))
    for sira, (ri, h) in enumerate(zip(idxler, proje.hammaddeler)):
        cells = t.rows[ri].cells
        bold = h.ara_toplam or (sira == son)
        # Ad (temiz) + yıldız
        ad = (h.ad or "").rstrip("*").strip()
        if getattr(h, "ucucu_sivi", False) and "ucucu" in yildiz_say:
            ad = ad + "*" * yildiz_say["ucucu"]
        elif ("kaplama materyali" in ad.lower() or getattr(h, "kaplama_yildiz", 0)) \
             and "kaplama" in yildiz_say:
            ad = ad + "*" * yildiz_say["kaplama"]
            h.kaplama_yildiz = yildiz_say["kaplama"]
        # Ad hücresi SOLA dayalı
        hucre_yaz(cells[0], ad, bold=bold, hiza="sol")
        hucre_yaz(cells[1], h.fonksiyon, bold=bold)
        # Uçucu sıvı: birim formül = k.m., % içerik = U.Y.
        if getattr(h, "ucucu_sivi", False):
            hucre_yaz(cells[2], "k.m.", bold=bold); var_km = True
            hucre_yaz(cells[3], "U.Y.", bold=bold); var_uy = True
        else:
            hucre_yaz(cells[2], _vf(h.birim_formul, 3), bold=bold)
            hucre_yaz(cells[3], _vf(h.yuzde_icerik, 2), bold=bold)
        # Seri Boyu: 3 hane + üst-karakter numara
        seri_txt = _vf(h.seri_miktar, 3)
        if h.ust_numara and h.ust_numara in _USTLER:
            seri_txt = seri_txt + _USTLER[h.ust_numara]
        hucre_yaz(cells[4], seri_txt, bold=bold)

    # --- Tablo altı notlar (bir kez) ---
    if not getattr(proje, "_formul_notlari_yazildi", False):
        # Şablonun SABİT (taslak) notlarını temizle (sadece ilk çağrıda)
        import re as _re2
        for p in list(doc.paragraphs):
            tx = (p.text or "").strip()
            if _re2.match(r"^\*+\s*:\s*Uçucudur", tx) or \
               _re2.match(r"^k\.m\.\s*:\s*kafi miktarda", tx) or \
               _re2.match(r"^U\.Y\.\s*:\s*Uygulama yoktur", tx):
                p._p.getparent().remove(p._p)
        _formul_notlari_ekle(doc, proje, _USTLER, tablo=t,
                             yildiz_say=yildiz_say, var_km=var_km, var_uy=var_uy)
        proje._formul_notlari_yazildi = True


def _formul_notlari_ekle(doc, proje, _USTLER, tablo=None, yildiz_say=None,
                         var_km=False, var_uy=False):
    """
    Birim formül tablosunun ALTINA notları ekler:
      - Uçucu sıvı notu (*/**): 'Uçucudur, birim ağırlığında yer almaz.'
      - Kaplama materyali bileşimi (*/**)
      - k.m.: kafi miktarda   (sadece tabloda k.m. varsa)
      - U.Y.: Uygulama yoktur (sadece tabloda U.Y. varsa)
      - Potens ve potens ayarlayıcı notları (¹, ²)
    Notlar tablonun hemen ardına eklenir.
    """
    yildiz_say = yildiz_say or {}
    seri_nolar = [s.seri_no for s in proje.seriler if s.seri_no]
    if len(seri_nolar) > 1:
        seri_metin = ", ".join(seri_nolar[:-1]) + " ve " + seri_nolar[-1]
    else:
        seri_metin = seri_nolar[0] if seri_nolar else ""

    anchor = tablo._tbl if tablo is not None else None
    olusturulan = []

    def _not_para(metin, ust=None):
        p = doc.add_paragraph()
        r = p.add_run((ust + " " if ust else "") + metin)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        olusturulan.append(p._p)
        return p

    # Uçucu sıvı notu (tek satır, yıldız sayısına göre)
    if "ucucu" in yildiz_say:
        y = "*" * yildiz_say["ucucu"]
        _not_para(f"{y}: Uçucudur, birim ağırlığında yer almaz.")

    # Kaplama materyali bileşimi notu (yıldız sayısına göre)
    for h in proje.hammaddeler:
        if getattr(h, "kaplama_yildiz", 0) and h.kaplama_bilesimi:
            yildiz = "*" * h.kaplama_yildiz
            _not_para(f"{yildiz}: Kaplama Materyali Bileşimi (% a/a) : {h.kaplama_bilesimi}")
            break

    # k.m. / U.Y. notları (sadece tabloda kullanıldıysa)
    if var_km:
        _not_para("k.m.: kafi miktarda")
    if var_uy:
        _not_para("U.Y.: Uygulama yoktur")

    # Potens notları (etken maddeler) + fazlalık toplamı
    fazlalik_toplam = 0.0
    for h in proje.hammaddeler:
        if getattr(h, "etken", False) and h.potens and h.seri_miktar:
            teorik = h.seri_miktar
            gercek = teorik * 100.0 / h.potens
            fazlalik_toplam += (gercek - teorik)
            ust = _USTLER.get(h.ust_numara or 0, "")
            _not_para(
                f": {seri_metin} serileri için teorik olarak " +
                f"{teorik:.2f}".replace(".", ",") +
                " kg tartılması gereken etkin madde miktarı, miktar ayarlaması sonrası " +
                f"{gercek:.2f}".replace(".", ",") + " kg olarak tartılmalıdır.", ust=ust)

    # Potens ayarlayıcı notu
    for h in proje.hammaddeler:
        if getattr(h, "potens_ayarlayici", False) and h.seri_miktar:
            teorik = h.seri_miktar
            gercek = teorik - fazlalik_toplam
            ust = _USTLER.get(h.ust_numara or 0, "")
            _not_para(
                f": {seri_metin} serileri için teorik olarak " +
                f"{teorik:.2f}".replace(".", ",") +
                f" kg tartılması gereken {h.ad} miktarı, miktar ayarlaması sonrası " +
                f"{gercek:.2f}".replace(".", ",") + " kg olarak tartılmalıdır.", ust=ust)

    # Notları tablonun hemen ardına taşı (sona değil)
    if anchor is not None and olusturulan:
        ref = anchor
        for p_el in olusturulan:
            p_el.getparent().remove(p_el)
            ref.addnext(p_el)
            ref = p_el


def _doldur_kapsanan(doc, proje: ProjeVerisi) -> None:
    t = _tablo_basliga_gore(doc, 2)  # Tablo 2
    if t is None:
        return
    # 2 başlık satırı (Ürün İsmi / Film Kaplı Tablet-kg alt başlığı)
    baslik = 2
    idxler = _veri_satirlarini_ayarla(t, baslik, SERI_SAYISI)
    for k, ri in enumerate(idxler):
        s = proje.seriler[k]
        cells = t.rows[ri].cells
        hucre_yaz(cells[0], s.urun_ismi or _urun(proje))
        hucre_yaz(cells[1], s.seri_no)
        hucre_yaz(cells[2], s.seri_boyutu_adet)
        hucre_yaz(cells[3], s.seri_boyutu_kg)


def _doldur_risk(doc, proje: ProjeVerisi) -> None:
    t = _tablo_basliga_gore(doc, 3)  # Tablo 3
    if t is None or not proje.risk_satirlari:
        return
    idxler = _veri_satirlarini_ayarla(t, 1, len(proje.risk_satirlari))
    for ri, rs in zip(idxler, proje.risk_satirlari):
        cells = t.rows[ri].cells
        hucre_yaz(cells[0], str(rs.operasyon_no or ""))
        hucre_yaz(cells[1], rs.operasyon)
        hucre_yaz(cells[2], "E" if rs.kritik else "H")
        hucre_yaz(cells[3], rs.testler)
        hucre_yaz(cells[4], rs.yorumlar)


def _doldur_proses_param(doc, proje: ProjeVerisi) -> None:
    t = _tablo_basliga_gore(doc, 4)  # Tablo 4
    if t is None:
        return
    # Kullanıcı öngörülen proses parametrelerini doldurmadıysa, şablondaki TASLAK
    # satırları (örn. 'Operasyon 2: Aşama 8 | Karıştırma Süresi | 10 dk') temizle.
    if not proje.proses_parametreleri:
        for ri in range(1, len(t.rows)):  # başlık satırı hariç
            for cell in t.rows[ri].cells:
                _hucre_temizle(cell)
        return
    idxler = _veri_satirlarini_ayarla(t, 1, len(proje.proses_parametreleri))
    for ri, pp in zip(idxler, proje.proses_parametreleri):
        cells = t.rows[ri].cells
        hucre_yaz(cells[0], pp.aciklama)
        hucre_yaz(cells[1], pp.parametre)
        if len(cells) > 2:
            hucre_yaz(cells[2], pp.deger)


def _hucre_temizle(cell):
    """Hücre içeriğini tamamen boşaltır (taslak metni siler)."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""


def _doldur_ekipman(doc, proje: ProjeVerisi) -> None:
    t = _tablo_basliga_gore(doc, 5)  # Tablo 5
    if t is None or not proje.ekipmanlar:
        return
    idxler = _veri_satirlarini_ayarla(t, 1, len(proje.ekipmanlar))
    for ri, e in zip(idxler, proje.ekipmanlar):
        cells = t.rows[ri].cells
        hucre_yaz(cells[0], str(e.operasyon_no or ""))
        hucre_yaz(cells[1], e.operasyon)
        hucre_yaz(cells[2], e.ekipman_adi)
        hucre_yaz(cells[3], e.kapasite)


def _doldur_numune(doc, proje: ProjeVerisi) -> None:
    t = _tablo_basliga_gore(doc, 10)  # Tablo 10
    if t is None or not proje.numune_plani:
        return
    idxler = _veri_satirlarini_ayarla(t, 1, len(proje.numune_plani))
    for ri, n in zip(idxler, proje.numune_plani):
        cells = t.rows[ri].cells
        hucre_yaz(cells[0], str(n.operasyon_no or ""))
        hucre_yaz(cells[1], n.operasyon)
        hucre_yaz(cells[2], n.numune_noktasi)
        hucre_yaz(cells[3], n.toplam_miktar)


def _doldur_spek(doc, proje: ProjeVerisi) -> None:
    """
    Tablo 6 — spesifikasyon. Türetilmiş test listesi tek geçişte yazılır.
    Alt-başlıklı testlerde (İlgili Bileşikler, Enantiomerik İmpurite, Mikrobiyolojik,
    Boyar Madde, Ağırlık Tekdüzeliği) grup başlığı satırında Op No + Operasyon dolu,
    ALT satırlarda BOŞ (dikey birleşik grup). Bulk + Tap Dansite tek satırda birleşir.
    """
    t = _tablo_basliga_gore(doc, 6)
    kart = proje.spek_karti
    if t is None or not kart.testler:
        return

    # Satır planı: (op_no, op, ad, spek)  — alt satırlarda op_no/op "" bırakılır
    plan = []
    _bulk_bekliyor = None  # Bulk Dansite satırını Tap ile birleştirmek için beklet
    for test in kart.testler:
        opno = str(test.operasyon_no or "")
        op = test.operasyon
        yildiz = "*" if test.yildizli else ""
        ad = test.ad + yildiz

        # --- Bulk + Tap Dansite → tek satır "Bulk ve Tap Dansite" ---
        _adl = _kucuk(test.ad)
        if "bulk dansite" in _adl:
            _bulk_bekliyor = (opno, op, test)
            continue
        if "tap dansite" in _adl:
            if _bulk_bekliyor is not None:
                bopno, bop, btest = _bulk_bekliyor
                spek = btest.spesifikasyon.metni_olustur() or "Bilgi amaçlıdır."
                plan.append((bopno, bop, "Bulk ve Tap Dansite" + yildiz, spek))
                _bulk_bekliyor = None
            else:
                plan.append((opno, op, "Bulk ve Tap Dansite" + yildiz,
                             test.spesifikasyon.metni_olustur() or "Bilgi amaçlıdır."))
            continue

        # ekstra alt satırlar (mikrobiyolojik, ağırlık tekdüzeliği, boyar)
        ekstra = list(test.alt_satirlar)
        if test.aciklama_etiketi:
            ekstra.append((test.aciklama_etiketi, test.aciklama_spek))
        if test.aciklama2_etiketi:
            ekstra.append((test.aciklama2_etiketi, test.aciklama2_spek))
        if ekstra:
            # başlık satırı (spek boş) op dolu; alt satırlarda op BOŞ
            plan.append((opno, op, ad, ""))
            for et, sp in ekstra:
                plan.append(("", "", et, sp))
        elif getattr(test, "_grup_baslik", False):
            # İlgili Bileşikler / Enantiomerik / Boyar başlığı: op dolu, spek boş
            plan.append((opno, op, ad, ""))
        elif getattr(test, "_impurite", False) or getattr(test, "_boyar_alt", False):
            # impurite/enantiomerik/boyar ALT satırı: op BOŞ (birleşik grup)
            plan.append(("", "", ad, test.spesifikasyon.spesifikasyon_metni
                         or test.spesifikasyon.metni_olustur() or ""))
        else:
            plan.append((opno, op, ad, test.spesifikasyon.metni_olustur()))

    idxler = _veri_satirlarini_ayarla(t, 1, len(plan))
    # Şablonun örnek satırlarında Op No/Operasyon hücrelerinde dikey birleştirme
    # (vMerge) olabilir; bu, boş yazılan alt-satır hücrelerinin üstteki değeri
    # devralmasına ve sütun kaymasına yol açar. Tüm tablodaki vMerge'leri kaldır.
    from docx.oxml.ns import qn as _qn
    for row in t.rows:
        for tc in row._tr.findall(_qn("w:tc")):
            tcPr = tc.find(_qn("w:tcPr"))
            if tcPr is not None:
                vm = tcPr.find(_qn("w:vMerge"))
                if vm is not None:
                    tcPr.remove(vm)
    for ri, (opno, op, ad, spek) in zip(idxler, plan):
        cells = t.rows[ri].cells
        hucre_yaz(cells[0], opno)
        hucre_yaz(cells[1], op)
        hucre_yaz(cells[2], ad)
        hucre_yaz(cells[3], spek)

    # --- Boş Op No / Operasyon hücrelerini üstteki grup başlığıyla BİRLEŞTİR ---
    # Ardışık (başlık + boş alt satırlar) bloklarını topla, her bloğu Op No ve
    # Operasyon sütunlarında TEK seferde dikey birleştir.
    bloklar = []  # (bas_ri, son_ri)
    bas = None
    for k, (opno, op, ad, spek) in enumerate(plan):
        ri = idxler[k]
        if opno.strip():
            if bas is not None and bas[1] > bas[0]:
                bloklar.append((bas[0], bas[1]))
            bas = [ri, ri]
        else:
            if bas is not None:
                bas[1] = ri
    if bas is not None and bas[1] > bas[0]:
        bloklar.append((bas[0], bas[1]))
    for bas_ri, son_ri in bloklar:
        try:
            t.rows[bas_ri].cells[0].merge(t.rows[son_ri].cells[0])
            t.rows[bas_ri].cells[1].merge(t.rows[son_ri].cells[1])
        except Exception:
            pass
    _tablo_genislik_duzelt(t)
    return


def _doldur_spek_ESKI(doc, proje: ProjeVerisi) -> None:
    """Tablo 6 — spesifikasyon. Yıldız test adının SONUNA; alt satırlar ayrı satır.
    İlgili Bileşikler etkin madde başına GRUPLU eklenir (operasyon hücresi dikey birleşik)."""
    t = _tablo_basliga_gore(doc, 6)  # Tablo 6
    kart = proje.spek_karti
    if t is None:
        return

    # 1) Normal testlerin satır planı
    satir_planı = []  # (tip, veri)
    for test in kart.testler:
        ekstra = list(test.alt_satirlar)
        if test.aciklama_etiketi:
            ekstra = ekstra + [(test.aciklama_etiketi, test.aciklama_spek)]
        if test.aciklama2_etiketi:
            ekstra = ekstra + [(test.aciklama2_etiketi, test.aciklama2_spek)]
        satir_planı.append(("test", (test, ekstra)))

    # 2) İlgili Bileşikler grupları (etkin madde başına)
    #    Her grup: başlık satırı + impurite satırları; operasyon dikey birleşik.
    ilgili_gruplar = []
    for em in kart.etkin_maddeler:
        if not em.impuriteler:
            continue
        ilk = em.impuriteler[0]
        yildiz = "*" if any(i.yildizli for i in em.impuriteler) else ""
        ilgili_gruplar.append({
            "baslik": f"{em.ad} İlgili Bileşikler{yildiz}",
            "operasyon": ilk.operasyon,
            "operasyon_no": ilk.operasyon_no,
            "impuriteler": em.impuriteler,
        })

    # Toplam satır sayısı
    toplam = 0
    for tip, veri in satir_planı:
        _, ekstra = veri
        toplam += 1 + len(ekstra)
    for g in ilgili_gruplar:
        toplam += 1 + len(g["impuriteler"])  # başlık + impurite satırları

    if toplam == 0:
        return

    idxler = _veri_satirlarini_ayarla(t, 1, toplam)
    it = iter(idxler)

    # Normal testler
    for tip, veri in satir_planı:
        test, ekstra = veri
        ri = next(it)
        cells = t.rows[ri].cells
        ad = test.ad + ("*" if test.yildizli else "")
        hucre_yaz(cells[0], str(test.operasyon_no or ""))
        hucre_yaz(cells[1], test.operasyon)
        hucre_yaz(cells[2], ad)
        ana_spek = "" if ekstra else test.spesifikasyon.metni_olustur()
        hucre_yaz(cells[3], ana_spek)
        for etiket, spek_metni in ekstra:
            ri2 = next(it)
            c2 = t.rows[ri2].cells
            hucre_yaz(c2[0], ""); hucre_yaz(c2[1], "")
            hucre_yaz(c2[2], etiket); hucre_yaz(c2[3], spek_metni)

    # İlgili Bileşikler grupları
    for g in ilgili_gruplar:
        grup_satir_idx = []
        # başlık satırı
        rb = next(it)
        grup_satir_idx.append(rb)
        cb = t.rows[rb].cells
        hucre_yaz(cb[0], str(g["operasyon_no"] or ""))
        hucre_yaz(cb[1], g["operasyon"])
        hucre_yaz(cb[2], g["baslik"])
        hucre_yaz(cb[3], "")
        # impurite satırları
        for imp in g["impuriteler"]:
            ri = next(it)
            grup_satir_idx.append(ri)
            c = t.rows[ri].cells
            hucre_yaz(c[0], ""); hucre_yaz(c[1], "")
            ad = imp.ad if imp.ad.startswith("—") or imp.ad.startswith("-") else f"—{imp.ad}"
            hucre_yaz(c[2], ad)
            hucre_yaz(c[3], imp.limit_metni or "")
        # Operasyon No + Operasyon hücrelerini grup boyunca DİKEY birleştir
        if len(grup_satir_idx) > 1:
            ust, alt = grup_satir_idx[0], grup_satir_idx[-1]
            t.rows[ust].cells[0].merge(t.rows[alt].cells[0])
            t.rows[ust].cells[1].merge(t.rows[alt].cells[1])
            hucre_yaz(t.rows[ust].cells[0], str(g["operasyon_no"] or ""))
            hucre_yaz(t.rows[ust].cells[1], g["operasyon"])


def _doldur_ipk(doc, proje: ProjeVerisi) -> None:
    """Tablo 7 — sadece IPK testleri. Ağırlık Tekdüzeliği başlık + 2 alt satır."""
    ipk = [t for t in proje.spek_karti.testler if t.ipk]
    t = _tablo_basliga_gore(doc, 7)  # Tablo 7
    if t is None or not ipk:
        return
    sut = len(t.columns)

    # Satır planı: her test 1 satır; Ağırlık Tekdüzeliği için + alt satırlar
    plan = []  # (op_no, op, ad, spek)
    for test in ipk:
        n = test.ad.lower()
        if "ağırlık tekdüzeliği" in n or "agirlik tekduzeligi" in n:
            plan.append((test.operasyon_no, test.operasyon, test.ad, ""))  # başlık boş
            if test.aciklama_etiketi:
                plan.append((test.operasyon_no, test.operasyon, test.aciklama_etiketi, test.aciklama_spek))
            if test.aciklama2_etiketi:
                plan.append((test.operasyon_no, test.operasyon, test.aciklama2_etiketi, test.aciklama2_spek))
        else:
            plan.append((test.operasyon_no, test.operasyon, test.ad, test.spesifikasyon.metni_olustur()))

    idxler = _veri_satirlarini_ayarla(t, 1, len(plan))
    for ri, (opno, op, ad, spek) in zip(idxler, plan):
        cells = t.rows[ri].cells
        if sut == 2:
            hucre_yaz(cells[0], ad)
            hucre_yaz(cells[1], spek)
        else:
            hucre_yaz(cells[0], str(opno or ""))
            hucre_yaz(cells[1], op)
            hucre_yaz(cells[2], ad)
            hucre_yaz(cells[3], spek)
    _tablo_genislik_duzelt(t)


# ============================================================================
# PVR sonuç tabloları (Bölüm 11) — şablonda örnek tablolar var; biz ekleyeceğiz
# ============================================================================
# Not: PVR sonuç tabloları çok sayıda ve teste göre değişken. Şablondaki örnek
# sonuç tablolarını silip, kullanıcı testlerine göre python-docx ile yeniden
# ekliyoruz (sonuç bölümü en sonda olduğu için bu güvenli).

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _sonuc_basligi(doc, no, ad):
    p = doc.add_paragraph()
    r = p.add_run(f"Tablo.{no} {ad} Sonuçları")
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    # Başlık daima takip eden tabloyla aynı sayfada kalsın
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    return p


def _yeni_tablo(doc, satir, sutun):
    t = doc.add_table(rows=satir, cols=sutun)
    t.style = "Table Grid"
    _tablo_bolunmesin(t)
    # Geniş tablolar (Baş/Orta/Son × 3 seri = 10 sütun) sayfaya sığsın
    if sutun >= SERI_SAYISI * 3 + 1:
        _genis_tablo_ayari(t)
    return t


def _tablo_bolunmesin(t):
    """
    Tablonun satırları sayfa sonunda bölünmesin (her satıra cantSplit) ve
    mümkün olduğunca tablo tek sayfada kalsın. Böylece bir tablo ikinci
    sayfaya sarkmaz.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)
    # Başlık satırını tekrar etme yok; sadece bölünmeyi engelliyoruz.


def _seri_dict(seriler, c):
    """seriler[c] güvenli dict erişimi (string/eksikse boş dict)."""
    if c < len(seriler) and isinstance(seriler[c], dict):
        return seriler[c]
    return {}


def _sr(cells, degerler, bold=False):
    for c, v in zip(cells, degerler):
        hucre_yaz(c, v, bold=bold) if c.paragraphs[0].runs else _yaz_bos(c, v, bold)


def _bicimle(metin):
    """Sayısal değerleri her zaman 2 ondalıkla biçimler (1 -> 1,00, 0.6 -> 0,60).
    Metin/None ise olduğu gibi döner. Türkçe ondalık ayracı (,) kullanılır."""
    if metin is None or metin == "":
        return ""
    if isinstance(metin, (int, float)):
        return f"{metin:.2f}".replace(".", ",")
    return str(metin)


def _yaz_bos(cell, metin, bold=False, punto=11):
    from docx.shared import RGBColor
    p = cell.paragraphs[0]
    r = p.add_run(_bicimle(metin))
    r.bold = bold
    r.font.size = Pt(punto)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Hücre içeriği tek satıra sığsın: paragraf boşluklarını sıfırla
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _yaz_sol(cell, metin, bold=False):
    """Sola dayalı hücre yazımı (Test/Spesifikasyon değer hücreleri için)."""
    from docx.shared import RGBColor
    p = cell.paragraphs[0]
    r = p.add_run(str(metin))
    r.bold = bold
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _numuneler_basligi(t, ust_satir, alt_satir, proje):
    """
    Resimlerdeki ortak yapı:
      | Numuneler | Analiz Sonuçları (3 sütuna birleşik)        |
      |           | Seri No: YYY-P01 | YYY-P02 | YYY-P03         |
    ust_satir: 'Numuneler' + 'Analiz Sonuçları' satırının indeksi
    alt_satir: Seri No satırının indeksi
    """
    _yaz_bos(t.rows[ust_satir].cells[0], "Numuneler", True)
    t.rows[ust_satir].cells[1].merge(t.rows[ust_satir].cells[SERI_SAYISI])
    _yaz_bos(t.rows[ust_satir].cells[1], "Analiz Sonuçları", True)
    # 'Numuneler' hücresini iki satıra dikey birleştir
    t.rows[ust_satir].cells[0].merge(t.rows[alt_satir].cells[0])
    for c, sno in enumerate(_seri_nolar(proje), 1):
        _yaz_bos(t.rows[alt_satir].cells[c], f"Seri No: {sno}", True)


def _ekle_sonuc_tek(doc, proje, test, no):
    """Görünüş/Teşhis/Elek: Test | Spesifikasyon | Numuneler+Analiz | Sonuç."""
    seriler = test.sonuc_verisi.get("seriler", ["", "", ""])
    _sonuc_basligi(doc, no, test.ad)
    t = _yeni_tablo(doc, 5, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur())
    _numuneler_basligi(t, 2, 3, proje)
    _yaz_bos(t.rows[4].cells[0], "Sonuç", True)
    for c in range(SERI_SAYISI):
        _yaz_bos(t.rows[4].cells[c+1], seriler[c] if c < len(seriler) else "", True)


def _ekle_sonuc_iki(doc, proje, test, no):
    """Miktar Tayini / İmpurite: Numune-1/Numune-2/Sonuç."""
    seriler = test.sonuc_verisi.get("seriler", [])
    _sonuc_basligi(doc, no, test.ad)
    t = _yeni_tablo(doc, 7, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur())
    _numuneler_basligi(t, 2, 3, proje)
    for ri, key, et in [(4, "numune_1", "Numune-1"), (5, "numune_2", "Numune-2"), (6, "sonuc", "Sonuç")]:
        _yaz_sol(t.rows[ri].cells[0], et, ri == 6)
        for c in range(SERI_SAYISI):
            v = _seri_dict(seriler, c).get(key, "")
            _yaz_bos(t.rows[ri].cells[c+1], v, ri == 6)
    # T.E. (tespit edilemedi) verisi varsa tablonun ALTINA not ekle
    if test.sonuc_verisi.get("te"):
        np = doc.add_paragraph()
        nr = np.add_run("T.E.: Tespit edilemedi.")
        nr.italic = True; nr.font.size = Pt(8)


def _ekle_sonuc_miktar(doc, proje, test, no):
    """
    Miktar Tayini (PVR Tablo.45/46): her seri Baş/Orta/Son; her nokta altında
    Numune-1, Numune-2, Sonuç. En altta seri Ortalaması (3 nokta birleşik).
    """
    seriler = test.sonuc_verisi.get("seriler", [])
    _sonuc_basligi(doc, no, test.ad)
    # Test + Spesifikasyon + Seri başlık + nokta başlık + N1 + N2 + Sonuç + Ortalama
    t = _yeni_tablo(doc, 8, SERI_SAYISI * 3 + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur())
    # Seri No (3'er birleşik) + nokta başlıkları
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    col = 1
    for sno in _seri_nolar(proje):
        a = t.rows[2].cells[col]; a.merge(t.rows[2].cells[col+2])
        _yaz_bos(a, f"Seri No: {sno}", True); col += 3
    t.rows[2].cells[0].merge(t.rows[3].cells[0])
    col = 1
    for _ in range(SERI_SAYISI):
        for nokta in NOKTA_ADLARI:
            _yaz_bos(t.rows[3].cells[col], nokta, True); col += 1
    # Numune-1 / Numune-2 / Sonuç satırları
    for ri, key, et in [(4, "numune_1", "Numune-1"), (5, "numune_2", "Numune-2"), (6, "sonuc", "Sonuç")]:
        _yaz_bos(t.rows[ri].cells[0], et, ri == 6)
        col = 1
        for c in range(SERI_SAYISI):
            noktalar = _seri_dict(seriler, c).get("noktalar", {})
            for nokta in NOKTA_ADLARI:
                _yaz_bos(t.rows[ri].cells[col], noktalar.get(nokta, {}).get(key, ""), ri == 6); col += 1
    # Ortalama satırı: her seri için tek değer (3 nokta birleşik)
    _yaz_bos(t.rows[7].cells[0], "Ortalama", True)
    col = 1
    for c in range(SERI_SAYISI):
        ort = _seri_dict(seriler, c).get("ortalama", "")
        a = t.rows[7].cells[col]; a.merge(t.rows[7].cells[col+2])
        _yaz_bos(a, ort, True); col += 3


def _genis_tablo_ayari(t):
    """
    Geniş (10 sütunlu) sonuç tablolarının sayfaya sığması için: otomatik
    genişlik (autofit) açılır ve tablo sayfa genişliğine yayılır. Değer
    hücrelerinde küçük punto (9) ile birlikte tek satıra sığma sağlanır.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    t.autofit = True
    t.allow_autofit = True
    tblPr = t._element.tblPr
    # Tablo genişliği: sayfanın tamamı (pct 5000 = %100)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")


def _ekle_sonuc_on(doc, proje, test, no):
    """Karışım Tekdüzeliği: 1-10 + Ortalama."""
    seriler = test.sonuc_verisi.get("seriler", [])
    _sonuc_basligi(doc, no, test.ad)
    t = _yeni_tablo(doc, 15, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur())
    _numuneler_basligi(t, 2, 3, proje)
    for n in range(10):
        _yaz_bos(t.rows[4+n].cells[0], str(n+1))
        for c in range(SERI_SAYISI):
            olc = _seri_dict(seriler, c).get("olcumler", [])
            _yaz_bos(t.rows[4+n].cells[c+1], olc[n] if n < len(olc) else "")
    _yaz_bos(t.rows[14].cells[0], "Ortalama", True)
    for c in range(SERI_SAYISI):
        _yaz_bos(t.rows[14].cells[c+1], _seri_dict(seriler, c).get("ortalama", ""), True)


def _ekle_sonuc_bos(doc, proje, test, no, ns=10):
    """Sertlik/Kalınlık/Çap/Dağılma/Dissolüsyon: Baş/Orta/Son × seri + Sonuç."""
    seriler = test.sonuc_verisi.get("seriler", [])
    _sonuc_basligi(doc, no, test.ad)
    # Test + Spesifikasyon + Seri başlık + nokta başlık + ns ölçüm + Ortalama + Sonuç
    t = _yeni_tablo(doc, 4 + ns + 2, SERI_SAYISI * 3 + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur())
    # seri başlıkları
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    col = 1
    for sno in _seri_nolar(proje):
        a = t.rows[2].cells[col]; a.merge(t.rows[2].cells[col+2])
        _yaz_bos(a, f"Seri No: {sno}", True); col += 3
    # 'Numuneler' dikey birleştir
    t.rows[2].cells[0].merge(t.rows[3].cells[0])
    col = 1
    for _ in range(SERI_SAYISI):
        for nokta in NOKTA_ADLARI:
            _yaz_bos(t.rows[3].cells[col], nokta, True); col += 1
    for n in range(ns):
        _yaz_bos(t.rows[4+n].cells[0], str(n+1))
        col = 1
        for c in range(SERI_SAYISI):
            for nokta in NOKTA_ADLARI:
                noktalar = _seri_dict(seriler, c).get("noktalar", {})
                olc = noktalar.get(nokta, {}).get("olcumler", [])
                _yaz_bos(t.rows[4+n].cells[col], olc[n] if n < len(olc) else ""); col += 1
    _yaz_bos(t.rows[4+ns].cells[0], "Ortalama", True)
    col = 1
    for c in range(SERI_SAYISI):
        for nokta in NOKTA_ADLARI:
            noktalar = _seri_dict(seriler, c).get("noktalar", {})
            _yaz_bos(t.rows[4+ns].cells[col], noktalar.get(nokta, {}).get("ortalama", ""), True); col += 1
    _yaz_bos(t.rows[5+ns].cells[0], "Sonuç", True)
    col = 1
    for c in range(SERI_SAYISI):
        sonuc = _seri_dict(seriler, c).get("sonuc", "")
        a = t.rows[5+ns].cells[col]; a.merge(t.rows[5+ns].cells[col+2])
        _yaz_bos(a, sonuc, True); col += 3
    _genis_tablo_ayari(t)


def _ekle_sonuc_ortalama_agirlik(doc, proje, test, no):
    """
    Ortalama Ağırlık. Kısma göre iki yapı:
      - Tablet: Seri No + Baş/Orta/Son; Sonuç (her nokta tek değer) + Ortalama.
      - Film  : Baş/Orta/Son YOK; her seri tek Sonuç değeri.
    Değerler Ağırlık Tekdüzeliği ortalamalarından gelir (birebir eşleşir).
    """
    if test.sonuc_verisi.get("film"):
        _ekle_sonuc_ortalama_agirlik_film(doc, proje, test, no)
        return
    seriler = test.sonuc_verisi.get("seriler", [])
    # Güvenlik: veri beklenen dict yapısında değilse (örn. test yanlış tiple
    # eklenmiş ve string listesi gelmiş) boş nokta yapısına çevir.
    def _nokta(c):
        if c < len(seriler) and isinstance(seriler[c], dict):
            return seriler[c].get("noktalar", {})
        return {}
    def _seri_sonuc(c):
        if c < len(seriler) and isinstance(seriler[c], dict):
            return seriler[c].get("sonuc", "")
        return ""
    _sonuc_basligi(doc, no, test.ad)
    # Test + Spesifikasyon + Numuneler + nokta başlık + Sonuç + Ortalama = 6 satır
    t = _yeni_tablo(doc, 6, SERI_SAYISI * 3 + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur())
    # Numuneler / Seri No (3'er birleşik) + nokta başlıkları
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    col = 1
    for sno in _seri_nolar(proje):
        a = t.rows[2].cells[col]; a.merge(t.rows[2].cells[col+2])
        _yaz_bos(a, f"Seri No: {sno}", True); col += 3
    t.rows[2].cells[0].merge(t.rows[3].cells[0])
    col = 1
    for _ in range(SERI_SAYISI):
        for nokta in NOKTA_ADLARI:
            _yaz_bos(t.rows[3].cells[col], nokta, True); col += 1
    # Sonuç satırı: her nokta tek değer (= Ağırlık Tekdüzeliği nokta ortalaması)
    _yaz_bos(t.rows[4].cells[0], "Sonuç", True)
    col = 1
    for c in range(SERI_SAYISI):
        noktalar = _nokta(c)
        for nokta in NOKTA_ADLARI:
            _yaz_bos(t.rows[4].cells[col], noktalar.get(nokta, {}).get("ortalama", ""), False); col += 1
    # Ortalama satırı: seri başına tek (3 nokta birleşik), seri ortalaması
    _yaz_bos(t.rows[5].cells[0], "Ortalama", True)
    col = 1
    for c in range(SERI_SAYISI):
        sonuc = _seri_sonuc(c)
        a = t.rows[5].cells[col]; a.merge(t.rows[5].cells[col+2])
        _yaz_bos(a, sonuc, True); col += 3


def _ekle_sonuc_ortalama_agirlik_film(doc, proje, test, no):
    """
    Ortalama Ağırlık — Film aşaması: Baş/Orta/Son YOK. Her seri tek sütun,
    tek Sonuç değeri (= o serinin 20 tabletinin ortalaması).
    """
    seriler = test.sonuc_verisi.get("seriler", [])
    def _ort(c):
        if c < len(seriler) and isinstance(seriler[c], dict):
            return seriler[c].get("ortalama", "")
        return ""
    _sonuc_basligi(doc, no, test.ad)
    # Test + Spesifikasyon + Numuneler(Seri No) + Sonuç
    t = _yeni_tablo(doc, 4, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur())
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    for c, sno in enumerate(_seri_nolar(proje)):
        _yaz_bos(t.rows[2].cells[c+1], f"Seri No: {sno}", True)
    _yaz_bos(t.rows[3].cells[0], "Sonuç", True)
    for c in range(SERI_SAYISI):
        _yaz_bos(t.rows[3].cells[c+1], _ort(c), True)


def _agirlik_spek_yaz(hucre, test):
    """
    Ağırlık Tekdüzeliği spesifikasyon hücresini PVR formatında doldurur:
    iki sapma satırı (etiket + limit). Sapma bilgisi (aciklama_etiketi/spek ve
    aciklama2_*) yoksa testin kendi spesifikasyon metnine düşer.
    """
    satirlar = []
    if test.aciklama_etiketi or test.aciklama_spek:
        et = (test.aciklama_etiketi or "").strip()
        sp = (test.aciklama_spek or "").strip()
        satirlar.append(f"{et} {sp}".strip())
    if test.aciklama2_etiketi or test.aciklama2_spek:
        et = (test.aciklama2_etiketi or "").strip()
        sp = (test.aciklama2_spek or "").strip()
        satirlar.append(f"{et} {sp}".strip())
    if not satirlar:
        _yaz_sol(hucre, test.spesifikasyon.metni_olustur())
        return
    # İlk satırı hücrenin mevcut paragrafına, kalanları yeni paragraf olarak yaz
    _yaz_sol(hucre, satirlar[0])
    for s in satirlar[1:]:
        p = hucre.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(s)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def _ekle_sonuc_agirlik(doc, proje, test, no):
    """
    Ağırlık Tekdüzeliği. Kısma göre iki yapı:
      - Tablet: Baş/Orta/Son × 10 numune + Ort/RSD%/SD.
      - Film  : Baş/Orta/Son YOK; seri başına düz 20 değer + Ort/RSD%/SD.
    Spesifikasyon hücresi iki sapma satırını gösterir (PVR Tablo.37/60).
    """
    if test.sonuc_verisi.get("film"):
        _ekle_sonuc_agirlik_film(doc, proje, test, no)
        return
    seriler = test.sonuc_verisi.get("seriler", [])
    _sonuc_basligi(doc, no, test.ad)
    t = _yeni_tablo(doc, 4 + 10 + 3, SERI_SAYISI * 3 + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI * 3])
    _agirlik_spek_yaz(t.rows[1].cells[1], test)
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    col = 1
    for sno in _seri_nolar(proje):
        a = t.rows[2].cells[col]; a.merge(t.rows[2].cells[col+2])
        _yaz_bos(a, f"Seri No: {sno}", True); col += 3
    t.rows[2].cells[0].merge(t.rows[3].cells[0])
    col = 1
    for _ in range(SERI_SAYISI):
        for nokta in NOKTA_ADLARI:
            _yaz_bos(t.rows[3].cells[col], nokta, True); col += 1
    for n in range(10):
        _yaz_bos(t.rows[4+n].cells[0], str(n+1))
        col = 1
        for c in range(SERI_SAYISI):
            for nokta in NOKTA_ADLARI:
                noktalar = _seri_dict(seriler, c).get("noktalar", {})
                olc = noktalar.get(nokta, {}).get("olcumler", [])
                _yaz_bos(t.rows[4+n].cells[col], olc[n] if n < len(olc) else ""); col += 1
    for k, (et, key) in enumerate([("Ortalama", "ortalama"), ("RSD%", "rsd"), ("SD", "sd")]):
        _yaz_bos(t.rows[14+k].cells[0], et, True)
        col = 1
        for c in range(SERI_SAYISI):
            for nokta in NOKTA_ADLARI:
                noktalar = _seri_dict(seriler, c).get("noktalar", {})
                _yaz_bos(t.rows[14+k].cells[col], noktalar.get(nokta, {}).get(key, ""), True); col += 1


def _ekle_sonuc_agirlik_film(doc, proje, test, no):
    """
    Ağırlık Tekdüzeliği — Film aşaması: Baş/Orta/Son YOK. Her seri tek sütun,
    seri başına 20 değer + Ortalama/RSD%/SD.
    """
    seriler = test.sonuc_verisi.get("seriler", [])
    _sonuc_basligi(doc, no, test.ad)
    # Test + Spesifikasyon + Seri başlık + 20 değer + Ort/RSD/SD
    t = _yeni_tablo(doc, 3 + 20 + 3, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _agirlik_spek_yaz(t.rows[1].cells[1], test)
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    for c, sno in enumerate(_seri_nolar(proje)):
        _yaz_bos(t.rows[2].cells[c+1], f"Seri No: {sno}", True)
    for n in range(20):
        _yaz_bos(t.rows[3+n].cells[0], str(n+1))
        for c in range(SERI_SAYISI):
            olc = _seri_dict(seriler, c).get("olcumler", [])
            _yaz_bos(t.rows[3+n].cells[c+1], olc[n] if n < len(olc) else "")
    for k, (et, key) in enumerate([("Ortalama", "ortalama"), ("RSD%", "rsd"), ("SD", "sd")]):
        _yaz_bos(t.rows[23+k].cells[0], et, True)
        for c in range(SERI_SAYISI):
            _yaz_bos(t.rows[23+k].cells[c+1], _seri_dict(seriler, c).get(key, ""), True)


def _ekle_sonuc_impurite(doc, proje, imp, baslik, no, tohum_ek=0, caption=None):
    """
    Bir impurite için sonuç tablosu:
      Test | <baslik>  (örn. 'Linezolid impurite C' veya 'Linezolid R-İzomer')
      Spesifikasyon | <limit metni>
      Numuneler/Analiz | Seri No
      Numune 1 / Numune 2 / Sonuç
    caption: tablo başlığı ('Tablo.X <caption> Sonuçları'). Verilmezse baslik kullanılır.
    Maksimum değere uyar; T.E. ise hepsi 'T.E.' + alt not.
    """
    import random as _r
    _sonuc_basligi(doc, no, caption if caption is not None else baslik)
    t = _yeni_tablo(doc, 7, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], baslik)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _yaz_sol(t.rows[1].cells[1], imp.limit_metni or "")
    # Numuneler / Analiz Sonuçları + Seri No
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    t.rows[2].cells[1].merge(t.rows[2].cells[SERI_SAYISI])
    _yaz_bos(t.rows[2].cells[1], "Analiz Sonuçları", True)
    t.rows[2].cells[0].merge(t.rows[3].cells[0])
    for c, sno in enumerate(_seri_nolar(proje), 1):
        _yaz_bos(t.rows[3].cells[c], f"Seri No: {sno}", True)

    te = imp.te or (str(imp.limit_metni).upper().replace(" ", "").endswith("T.E.")
                    if imp.limit_metni else False)
    maks = imp.maksimum_deger
    # KULLANICI KURALI: impurite değerleri asla 0.07'den yüksek olmaz ve spek
    # üst sınırını da aşmaz.
    ust = 0.07
    if maks:
        ust = min(ust, maks)
    # Satır etiketleri
    _yaz_sol(t.rows[4].cells[0], "Numune 1")
    _yaz_sol(t.rows[5].cells[0], "Numune 2")
    _yaz_sol(t.rows[6].cells[0], "Sonuç", True)
    # Numune 1/2 üret, Sonuç = ortalama; T.E. ise hepsi T.E.
    for c in range(SERI_SAYISI):
        if te:
            n1 = n2 = s = "T.E."
        else:
            v1 = round(_r.uniform(0.0, ust), 2)
            v2 = round(_r.uniform(0.0, ust), 2)
            n1, n2 = _bicimle_sayi(v1), _bicimle_sayi(v2)
            s = _bicimle_sayi(round((v1 + v2) / 2, 2))
        _yaz_bos(t.rows[4].cells[c+1], n1)
        _yaz_bos(t.rows[5].cells[c+1], n2)
        _yaz_bos(t.rows[6].cells[c+1], s, True)
    if te:
        np = doc.add_paragraph()
        nr = np.add_run("T.E.: Tespit edilemedi.")
        nr.italic = True; nr.font.size = Pt(8)


def _bicimle_sayi(deger, ondalik=2) -> str:
    """Çıktı motoru içi: sayıyı 2 ondalık + Türkçe virgülle string yapar."""
    try:
        return f"{float(deger):.{ondalik}f}".replace(".", ",")
    except (ValueError, TypeError):
        return str(deger)


def _ekle_limit_tablosu(doc, proje, baslik, limit_metni, no):
    """
    Ağırlık Tekdüzeliği limit tablosu (resim 2 — Tablo 35/36):
      Test | <başlık>
      Spesifikasyon | <limit metni>
      Numuneler / Analiz Sonuçları | Seri No
      Sonuç | <limit metni> × 3 seri
    """
    doc.add_paragraph("")
    _sonuc_basligi(doc, no, baslik)
    t = _yeni_tablo(doc, 5, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], baslik)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _yaz_sol(t.rows[1].cells[1], limit_metni)
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    t.rows[2].cells[1].merge(t.rows[2].cells[SERI_SAYISI])
    _yaz_bos(t.rows[2].cells[1], "Analiz Sonuçları", True)
    t.rows[2].cells[0].merge(t.rows[3].cells[0])
    for c, sno in enumerate(_seri_nolar(proje), 1):
        _yaz_bos(t.rows[3].cells[c], f"Seri No: {sno}", True)
    _yaz_bos(t.rows[4].cells[0], "Sonuç", True)
    for c in range(SERI_SAYISI):
        _yaz_bos(t.rows[4].cells[c+1], limit_metni, False)


def _ekle_sonuc_sizdirmazlik(doc, proje, test, no):
    """
    Sızdırmazlık (şablon Tablo.77): Baş/Orta/Son satırları, her seri 3 alt sütun
    (1/2/3), tüm hücreler 'Uygun'; en altta Sonuç satırı (seri başına 'Uygun').
    """
    _sonuc_basligi(doc, no, test.ad)
    # Test + Spesifikasyon + Numuneler(SeriNo) + altbaşlık(1/2/3) + Baş/Orta/Son + Sonuç
    t = _yeni_tablo(doc, 4 + 3 + 1, SERI_SAYISI * 3 + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[0].cells[1], test.ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI * 3])
    _yaz_sol(t.rows[1].cells[1], test.spesifikasyon.metni_olustur() or "Sızdırmamalı")
    # Numuneler / Seri No (3'er birleşik)
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    col = 1
    for sno in _seri_nolar(proje):
        a = t.rows[2].cells[col]; a.merge(t.rows[2].cells[col+2])
        _yaz_bos(a, f"Seri No: {sno}", True); col += 3
    # alt başlık: 1/2/3
    t.rows[2].cells[0].merge(t.rows[3].cells[0])
    col = 1
    for _ in range(SERI_SAYISI):
        for n in ("1", "2", "3"):
            _yaz_bos(t.rows[3].cells[col], n, True); col += 1
    # Baş/Orta/Son satırları → hepsi 'Uygun'
    for ri, et in [(4, "Baş"), (5, "Orta"), (6, "Son")]:
        _yaz_bos(t.rows[ri].cells[0], et, True)
        for c in range(1, SERI_SAYISI * 3 + 1):
            _yaz_bos(t.rows[ri].cells[c], "Uygun")
    # Sonuç satırı: her seri (3 birleşik) 'Uygun'
    _yaz_bos(t.rows[7].cells[0], "Sonuç", True)
    col = 1
    for _ in range(SERI_SAYISI):
        a = t.rows[7].cells[col]; a.merge(t.rows[7].cells[col+2])
        _yaz_bos(a, "Uygun", True); col += 3


def _ekle_sonuc_matris(doc, proje, test, no):
    """
    Mikrobiyolojik Kontrol (resimdeki Tablo 33/34):
      Test | Mikrobiyolojik Kontrol
      Spesifikasyon | (3 alt satır spek, birleşik)
      Numuneler/Analiz Sonuçları | Seri No
      -Toplam Aerobik... | <10 cfu/g × 3
      -Küf ve Maya...    | <10 cfu/g × 3
      -E. coli           | 0 cfu/g × 3
    Sonuç satırı YOK (kullanıcı kararı).
    """
    _sonuc_basligi(doc, no, test.ad)
    alt = test.alt_satirlar or [("-Toplam Aerobik Mikroorganizma Sayısı", "≤10³ cfu/g"),
                                ("-Küf ve Maya Sayısı", "≤10² cfu/g"), ("-E. coli", "0 cfu/g")]
    # Test(1) + Spesifikasyon(len alt, birleşik) + Numuneler/SeriNo(2) + ölçüm(len alt)
    n_spec = len(alt)
    t = _yeni_tablo(doc, 1 + n_spec + 2 + len(alt), SERI_SAYISI + 1)
    r = 0
    _yaz_bos(t.rows[r].cells[0], "Test", True)
    t.rows[r].cells[1].merge(t.rows[r].cells[SERI_SAYISI])
    _yaz_sol(t.rows[r].cells[1], test.ad); r += 1
    # Spesifikasyon: sol hücre dikey birleşik, sağda alt satır spek'leri
    spec_bas = r
    for i, (etiket, spek_metni) in enumerate(alt):
        t.rows[r].cells[1].merge(t.rows[r].cells[SERI_SAYISI])
        _yaz_sol(t.rows[r].cells[1], f"{etiket} : {spek_metni}")
        r += 1
    _yaz_bos(t.rows[spec_bas].cells[0], "Spesifikasyon", True)
    t.rows[spec_bas].cells[0].merge(t.rows[r-1].cells[0])
    # Numuneler / Analiz Sonuçları + Seri No
    _numuneler_basligi(t, r, r+1, proje); r += 2
    # ölçüm satırları
    for etiket, spek_metni in alt:
        _yaz_sol(t.rows[r].cells[0], etiket)
        for c in range(SERI_SAYISI):
            # değer: aerobik/küf <10 cfu/g, e.coli 0 cfu/g
            deger = "0 cfu/g" if "coli" in etiket.lower() else "<10 cfu/g"
            _yaz_bos(t.rows[r].cells[c+1], deger)
        r += 1


def _genel_degerlendirme_paragrafi(doc):
    """Şablondaki 'GENEL DEĞERLENDİRME' başlık paragrafını bulur (varsa)."""
    for p in doc.paragraphs:
        if "GENEL DEĞERLENDİRME" in p.text.strip().upper():
            return p
    return None


def _tum_belge_fontu(doc, font_adi: str = "Times New Roman") -> None:
    """
    Tüm belgedeki yazı tipini Times New Roman yapar: Normal stil + tüm
    paragraf run'ları + tüm tablo hücreleri. East Asian font da ayarlanır ki
    Word her durumda TNR uygulasın.
    """
    from docx.oxml.ns import qn
    # Normal stil
    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_adi
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font_adi)
    except Exception:
        pass

    def _run_font(run):
        run.font.name = font_adi
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font_adi)

    for p in doc.paragraphs:
        for r in p.runs:
            _run_font(r)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _run_font(r)


def _operasyon_basligi(operasyon: str) -> str:
    """Operasyon adını sonuç bölümü başlığına çevirir."""
    op = _kucuk((operasyon or "").strip())
    if not op:
        return ""
    if "karış" in op or "karis" in op:
        return "Karışım Aşaması"
    if "tablet" in op:
        return "Tablet Baskı Aşaması"
    if "film" in op:
        return "Film Kaplama Aşaması"
    if "blister" in op or "ambalaj" in op:
        return "Blisterleme Aşaması"
    return f"{operasyon} Aşaması"


def _operasyon_basligi_yaz(doc, baslik: str):
    """Sonuç bölümünde aşama başlığını (kalın) yazar."""
    p = doc.add_paragraph()
    r = p.add_run(baslik)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"
    p.paragraph_format.keep_with_next = True


def _doldur_genel_degerlendirme(doc, proje: ProjeVerisi) -> None:
    """
    Genel Değerlendirme tablolarını (Sapmalar / Sonuçlar ve Değerlendirme / Yorum)
    sabit standart metinlerle doldurur. Bu metinler her PVR'de aynıdır.
    """
    urun = _urun(proje).upper() if hasattr(_urun(proje), "upper") else str(_urun(proje))
    yorum = ("Yapılan üretimlerin validasyonları tamamlandıktan sonra hazırlanmış "
             f"olan rapor üretilecek olan {urun}'in sürekli olarak aynı "
             "spesifikasyonlarda üretiminin sağlanabileceğini kanıtlamaktadır.")
    for t in doc.tables:
        baslik = _kucuk((t.rows[0].cells[0].text or "").strip())
        if baslik.startswith("sapmalar") and len(t.rows) >= 2:
            if not t.rows[1].cells[0].text.strip():
                _hucre_kalin_yaz(t.rows[1].cells[0], "Sapma gözlenmemiştir.")
                _hucre_kalin_yaz(t.rows[1].cells[1], "U.Y.")
        elif baslik.startswith("sonuçlar ve değerlendirme") and len(t.rows) >= 2:
            if not t.rows[1].cells[0].text.strip():
                _hucre_kalin_yaz(t.rows[1].cells[0],
                                 "Belirtilen spesifikasyonlara uygun şekilde sonuçlar elde edilmiştir.")
        elif baslik.startswith("yorum") and len(t.rows) >= 2:
            if not t.rows[1].cells[0].text.strip():
                _hucre_kalin_yaz(t.rows[1].cells[0], yorum)

    # Genel Değerlendirme bölümü TEK sayfada ve başlık sayfanın EN ÜSTÜNDE olsun:
    # başlık paragrafına 'sayfa öncesi kesme' (page-break-before) ekle; tabloların
    # satırlarını bölünmez yap (cantSplit) böylece bölüm alt sayfaya sarkmaz.
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Ox
    for p in doc.paragraphs:
        if _kucuk(p.text.strip()).startswith("genel değerlendirme") or \
           _kucuk(p.text.strip()).startswith("genel degerlendirme"):
            pPr = p._p.get_or_add_pPr()
            pb = pPr.find(_qn("w:pageBreakBefore"))
            if pb is None:
                pb = _Ox("w:pageBreakBefore"); pPr.append(pb)
            kn = _Ox("w:keepNext"); pPr.append(kn)
            break
    # Genel Değerlendirme tablolarının satırları bölünmesin
    for t in doc.tables:
        b = _kucuk((t.rows[0].cells[0].text or "").strip())
        if b.startswith(("sapmalar", "sonuçlar ve değerlendirme", "yorum")):
            for row in t.rows:
                trPr = row._tr.get_or_add_trPr()
                trPr.append(_Ox("w:cantSplit"))


def _hucre_kalin_yaz(cell, metin):
    """Genel Değerlendirme hücresine kalın, Times New Roman metin yazar."""
    from docx.shared import RGBColor
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(metin)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 0)


def _ekle_sonuc_boyar(doc, proje, test_ad, spek_metni, no):
    """Boyar Madde alt testi için sonuç tablosu (Test | ad, Spesifikasyon,
    Numuneler/Seri No, Sonuç=Pozitif × seri). Tüm tablolar aynı yapıda."""
    _sonuc_basligi(doc, no, "Boyar Madde Tanıması")
    t = _yeni_tablo(doc, 4, SERI_SAYISI + 1)
    _yaz_bos(t.rows[0].cells[0], "Test", True)
    t.rows[0].cells[1].merge(t.rows[0].cells[SERI_SAYISI])
    _yaz_sol(t.rows[0].cells[1], test_ad)
    _yaz_bos(t.rows[1].cells[0], "Spesifikasyon", True)
    t.rows[1].cells[1].merge(t.rows[1].cells[SERI_SAYISI])
    _yaz_sol(t.rows[1].cells[1], spek_metni or "")
    _yaz_bos(t.rows[2].cells[0], "Numuneler", True)
    for c, sno in enumerate(_seri_nolar(proje), 1):
        _yaz_bos(t.rows[2].cells[c], f"Seri No: {sno}", True)
    _yaz_bos(t.rows[3].cells[0], "Sonuç", True)
    for c in range(1, SERI_SAYISI + 1):
        _yaz_bos(t.rows[3].cells[c], "Pozitif", True)


def _test_to_impurite(test):
    """Türetilmiş impurite/enantiomerik alt-satır Test'inden, _ekle_sonuc_impurite'ın
    beklediği hafif bir impurite nesnesi üretir."""
    from core.models import Impurite
    sp = test.spesifikasyon
    metin = sp.spesifikasyon_metni or sp.metni_olustur() or ""
    return Impurite(
        ad=test.ad.lstrip("—-– ").strip(),
        limit_metni=metin,
        maksimum_deger=sp.maksimum_deger,
        te=(str(sp.maksimum_metin or "").upper().replace(" ", "").endswith("T.E.")),
    )


def _doldur_sonuclar(doc, proje: ProjeVerisi) -> None:
    """
    PVR sonuç tablolarını şablondaki 'GENEL DEĞERLENDİRME' başlığının ÖNÜNE ekler.
    Böylece Genel Değerlendirme HER ZAMAN en sonda kalır.
    """
    hedef = _genel_degerlendirme_paragrafi(doc)

    # Sonuç bloğunu belge SONUNA üret (geçici), sonra hedefin önüne taşı.
    # Üretilen yeni elemanları belirlemek için mevcut body çocuklarını işaretle.
    from docx.oxml.ns import qn
    body = doc.element.body
    onceki = list(body)

    # başlık
    h = doc.add_paragraph()
    r = h.add_run("PROSES VALİDASYONU TEST SONUÇLARI")
    r.bold = True; r.font.size = Pt(12)

    no = 11
    onceki_op = None
    for test in proje.spek_karti.testler:
        tip = test.tablo_tipi
        ad_l = _kucuk(test.ad)

        # Grup başlıkları (İlgili Bileşikler / Enantiomerik) tablo üretmez.
        if getattr(test, "_grup_baslik", False) and \
           ("ilgili" in ad_l or "enantiomerik" in ad_l):
            continue
        if ("ilgili bileşik" in ad_l or "ilgili bilesik" in ad_l) and \
           tip in (TabloTipi.TEK_SONUC, TabloTipi.IKI_NUMUNE) and \
           not getattr(test, "_impurite", False):
            continue
        # Boyar Madde alt satırı tek başına burada işlenmez (grup başlığında ele alınır).
        if getattr(test, "_boyar_alt", False):
            continue

        # Operasyon başlığı: aşama değiştiğinde okunabilir başlık yaz
        op_baslik = _operasyon_basligi(test.operasyon)
        if op_baslik and op_baslik != onceki_op:
            _operasyon_basligi_yaz(doc, op_baslik)
            onceki_op = op_baslik

        # Boyar Madde: alt başlıklara göre sonuç tablosu/tabloları.
        #  - Tek alt test (örn. Titanyum dioksit): tek tablo, Test="Boyar Madde
        #    Tanıması", Spesifikasyon=alt test speki.
        #  - Çoklu alt test: her alt test için AYRI tablo (aynı yapı).
        if getattr(test, "_boyar", False) and getattr(test, "_grup_baslik", False):
            altlar = list(getattr(test, "_alt_basliklar", []) or [])
            if len(altlar) <= 1:
                sp = altlar[0][1] if altlar else (test.spesifikasyon.spesifikasyon_metni or "")
                _ekle_sonuc_boyar(doc, proje, "Boyar Madde Tanıması", sp, no)
                doc.add_paragraph(""); no += 1
            else:
                for alt_ad, alt_sp in altlar:
                    _ekle_sonuc_boyar(doc, proje, alt_ad.lstrip("—-– ").strip(), alt_sp, no)
                    doc.add_paragraph(""); no += 1
            continue

        # İlgili Bileşikler / Enantiomerik ALT satırı → o aşamada sonuç tablosu
        if getattr(test, "_impurite", False) or getattr(test, "_enantiomerik", False):
            imp = _test_to_impurite(test)
            test_ad = test.ad.lstrip("—-– ").strip()
            cap = "Enantiomerik İmpurite" if getattr(test, "_enantiomerik", False) else "İlgili Bileşikler"
            tohum_ek = 7 if getattr(test, "_enantiomerik", False) else 0
            _ekle_sonuc_impurite(doc, proje, imp, test_ad, no, tohum_ek=tohum_ek, caption=cap)
            doc.add_paragraph("")
            no += 1
            continue

        if test.mikrobiyolojik or tip is TabloTipi.MATRIS:
            _ekle_sonuc_matris(doc, proje, test, no)
        elif "sızdırmazlık" in ad_l or "sizdirmazlik" in ad_l:
            _ekle_sonuc_sizdirmazlik(doc, proje, test, no)
        elif "ortalama ağırlık" in ad_l:
            _ekle_sonuc_ortalama_agirlik(doc, proje, test, no)
        elif tip is TabloTipi.TEK_SONUC:
            _ekle_sonuc_tek(doc, proje, test, no)
        elif tip is TabloTipi.IKI_NUMUNE:
            if "miktar tayini" in ad_l:
                _ekle_sonuc_miktar(doc, proje, test, no)
            else:
                _ekle_sonuc_iki(doc, proje, test, no)
        elif tip is TabloTipi.ON_NUMUNE:
            _ekle_sonuc_on(doc, proje, test, no)
        elif tip is TabloTipi.BOS_NOKTA:
            if "dissol" in ad_l:
                _ekle_sonuc_bos(doc, proje, test, no, ns=6)
            else:
                _ekle_sonuc_bos(doc, proje, test, no)
        elif tip is TabloTipi.AGIRLIK_TEKDUZELIGI:
            _ekle_sonuc_agirlik(doc, proje, test, no)
        else:
            _ekle_sonuc_tek(doc, proje, test, no)
        doc.add_paragraph("")
        no += 1

    # NOT: İlgili Bileşikler ve Enantiomerik İmpurite sonuç tabloları artık
    # dispatcher içinde HER AŞAMADA (karışım/tablet/film) üretiliyor; burada
    # ayrıca üretilmez (çift basımı önlemek için).

    # Yeni eklenen elemanlar (başlık + tablolar + boş paragraflar)
    yeni_elemanlar = [el for el in body if el not in onceki]

    if hedef is not None:
        hedef_el = hedef._p
        for el in yeni_elemanlar:
            body.remove(el)
            hedef_el.addprevious(el)
    # hedef yoksa zaten en sonda kalır (sorun değil)


# ============================================================================
# Ana üretim
# ============================================================================

def _revizyon_no_guncelle(doc, revizyon_no: str) -> None:
    """
    Header/footer'daki Revizyon No değerini kullanıcının değeriyle günceller.
    İki yerleşim de desteklenir:
      - Aynı hücrede etiket+değer ('Revizyon No\\nNumber of Revision\\n00')
      - Etiket üst satırda, değer alt satırda ayrı hücrede (footer)
    """
    if not revizyon_no:
        return
    rev = str(revizyon_no).strip()

    def _son_sayiyi_degistir(cell):
        """Hücredeki saf-sayı paragrafını/run'ını rev ile değiştirir."""
        for p in cell.paragraphs:
            # tüm paragraf metni saf sayı mı?
            if p.text.strip().isdigit() and p.text.strip() != rev:
                # ilk run'a yaz, kalanları temizle
                if p.runs:
                    p.runs[0].text = rev
                    for r in p.runs[1:]:
                        r.text = ""
                continue
            for r in p.runs:
                t = r.text.strip()
                if t.isdigit() and 1 <= len(t) <= 3 and t != rev:
                    r.text = r.text.replace(t, rev)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for t in hf.tables:
                ncols = len(t.columns)
                # Etiket satırında 'Revizyon No' olan SÜTUNU bul
                rev_col = None
                for ri, row in enumerate(t.rows):
                    for ci, cell in enumerate(row.cells):
                        if "revizyon no" in _kucuk(cell.text) or "revision no" in _kucuk(cell.text):
                            # Aynı hücrede değer de varsa düzelt
                            _son_sayiyi_degistir(cell)
                            rev_col = ci
                # Footer tipi: etiket üstte, değer altta → o sütunun diğer satırı
                if rev_col is not None and len(t.rows) >= 2:
                    for ri, row in enumerate(t.rows):
                        cell = row.cells[rev_col]
                        if "revizyon" not in _kucuk(cell.text):  # etiket hücresi değil
                            _son_sayiyi_degistir(cell)


def _revizyon_tarihi_guncelle(doc, tarih: str) -> None:
    """Header/footer'daki 'Revizyon Tarihi' hücresindeki değeri (U.Y. veya tarih)
    kullanıcının girdiği tarihle değiştirir."""
    if not tarih:
        return
    tarih = str(tarih).strip()
    import re as _re
    # Değer deseni: gg.aa.yyyy veya U.Y.
    def _deger_degistir(cell):
        for p in cell.paragraphs:
            for r in p.runs:
                t = r.text.strip()
                if _re.match(r"^\d{2}\.\d{2}\.\d{4}$", t) or t in ("U.Y.", "U.Y", "UY"):
                    if t != tarih:
                        r.text = r.text.replace(t, tarih)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for t in hf.tables:
                tar_col = None
                for row in t.rows:
                    for ci, cell in enumerate(row.cells):
                        if "revizyon tarih" in _kucuk(cell.text) or "date of revision" in _kucuk(cell.text):
                            _deger_degistir(cell)
                            tar_col = ci
                if tar_col is not None and len(t.rows) >= 2:
                    for row in t.rows:
                        cell = row.cells[tar_col]
                        if "revizyon" not in _kucuk(cell.text) and "date" not in _kucuk(cell.text):
                            _deger_degistir(cell)


def _placeholder_eslemeleri(proje: ProjeVerisi, rapor: bool) -> dict[str, str]:
    d = proje.dokuman
    urun = _urun(proje)
    dok_no = (d.pvr_dokuman_no if rapor else d.pvp_dokuman_no) or "AG-PV-xxx"
    # Form No: PVR ise pvr_form_no, PVP ise pvp_form_no (kullanıcı değiştirebilir)
    form_no = (d.pvr_form_no if rapor else d.pvp_form_no) or ("N-15-507" if rapor else "N-15-506")
    es = {
        "XxxFilm Kaplı Tablet": urun,
        "XxxFİLM TABLET": urun.upper(),
        "AG-PV-xxx": dok_no,
        # Şablon footer'ındaki sabit form no'yu kullanıcının değeriyle değiştir
        "N-15-0506": form_no,
        "N-15-506": form_no,
        "N-15-507": form_no,
    }
    for i in range(SERI_SAYISI):
        sno = proje.seriler[i].seri_no
        if sno:
            es[f"yyy-P0{i+1}"] = sno

    # Amaç metnindeki firma ifadesi: '{Firma ismi} İlaç üretim fabrikasında' →
    # 'Neutec İlaç üretim fabrikasında' (çift 'İlaç' oluşmasını önle).
    es["{Firma ismi} İlaç üretim fabrikasında"] = "Neutec İlaç üretim fabrikasında"
    es["{Firma ismi} İlaç"] = "Neutec İlaç"
    if d.firma_ismi:
        es["{Firma ismi}"] = d.firma_ismi

    # --- Kapsam / Sorumluluk / Kaydedilme / Amaç metinleri (PVP vs PVR tense) ---
    es["sorumluluk dağılımları sayfa 5’teki gibidir."] = \
        "sorumluluk dağılımları aşağıdaki gibidir."
    es["sorumluluk dağılımları sayfa 5'teki gibidir."] = \
        "sorumluluk dağılımları aşağıdaki gibidir."
    if rapor:  # PVR → geçmiş zaman
        es["validasyon çalışması ardışık 3 seriye uygulanacaktır."] = \
            "validasyon çalışması ardışık 3 seriye uygulanmıştır."
        es["Raporunda verilen tablolara kaydedilecektir."] = \
            "Raporunda verilen tablolara kaydedilmiştir."
        es["tablolara kaydedilecektir."] = "tablolara kaydedilmiştir."
        es["üretim fabrikasında gerçekleştirilecektir."] = \
            "üretim fabrikasında gerçekleştirilmiştir."
    return es


def _tolerans_oran(tolerans: str):
    """'±%5' / '%7.5' / '±%10' gibi metinden oranı (0.05) çıkarır."""
    import re
    m = re.search(r"(\d+(?:[.,]\d+)?)", tolerans or "")
    if not m:
        return None
    return float(m.group(1).replace(",", ".")) / 100.0


def _miktar_spek_uret(hedef: float, tolerans: str, birim: str) -> str:
    """Hedef + tolerans → '10.0 mg/f.tab ±%7.5 (9.25 – 10.75 mg/f.tab)'."""
    oran = _tolerans_oran(tolerans)
    birim = (birim or "").strip()
    if oran is None:
        return f"{hedef:g} {birim}".strip()
    alt = round(hedef * (1 - oran), 4)
    ust = round(hedef * (1 + oran), 4)
    if birim:
        return f"{hedef:g} {birim} {tolerans} ({alt:g} – {ust:g} {birim})"
    return f"{hedef:g} {tolerans} ({alt:g} – {ust:g})"


def _test_bul(testler, *anahtarlar):
    """Adında verilen anahtar kelimeleri içeren ilk testi bulur."""
    for t in testler:
        ad = t.ad.lower()
        if all(a.lower() in ad for a in anahtarlar):
            return t
    return None


def _doldur_tablo89(doc, proje: ProjeVerisi) -> None:
    """
    Tablo 8/9 doldurucu dispatcher. Kullanıcı Word'den tablo yüklediyse ham XML
    ile BİREBİR kopyalanır (biçim korunur); değilse eski (yeniden inşa) yöntem.
    """
    kart = proje.spek_karti
    ham = getattr(kart, "_tablo8_xml", None)
    if ham is not None and _tablo89_birebir_kopya(doc, proje, ham):
        return
    _doldur_tablo89_eski(doc, proje)


def _tablo89_birebir_kopya(doc, proje: ProjeVerisi, ham_tbl) -> bool:
    """
    Şablondaki Tablo 8'i (ve varsa Tablo 9'u), kullanıcının yüklediği Word
    tablosunun ham XML'iyle BİREBİR değiştirir (biçim: kalın/italik/boyut korunur).
    Tablo 9 = Tablo 8 kopyası + Miktar Tayini satırı raf ömrü toleransıyla düzenli.
    Başarılıysa True döner.
    """
    import copy as _copy
    from docx.oxml.ns import qn
    kart = proje.spek_karti

    def _sablon_tablo(no):
        return _tablo_basliga_gore(doc, no)

    t8 = _sablon_tablo(8)
    if t8 is None:
        return False

    # --- Tablo 8: şablon tablosunu ham XML ile değiştir ---
    yeni8 = _copy.deepcopy(ham_tbl)
    t8._tbl.addprevious(yeni8)
    t8._tbl.getparent().remove(t8._tbl)

    # --- Tablo 9: kopya + Miktar Tayini raf ömrü toleransı ---
    raf_ekle = getattr(kart, "tablo89_ekle", True)
    t9 = _sablon_tablo(9)
    if t9 is not None:
        if raf_ekle:
            yeni9 = _copy.deepcopy(ham_tbl)
            # Miktar Tayini hücresini raf ömrü toleransıyla güncelle
            _tablo9_miktar_guncelle(yeni9, kart)
            t9._tbl.addprevious(yeni9)
            t9._tbl.getparent().remove(t9._tbl)
        else:
            # Tablo 9 istenmiyorsa şablon taslağını temizle
            for ri in range(1, len(t9.rows)):
                for cell in t9.rows[ri].cells:
                    _hucre_temizle(cell)
    return True


def _tablo9_miktar_guncelle(tbl_el, kart) -> None:
    """Kopyalanan Tablo 9 XML'inde Miktar Tayini satırının spesifikasyonunu
    raf ömrü toleransıyla (varsayılan ±%10) yeniden yazar."""
    from docx.table import Table
    from docx.oxml.ns import qn
    # tbl_el bir CT_Tbl; Table sarmalayıcı ile satırlara eriş
    try:
        t = Table(tbl_el, None)
    except Exception:
        return
    tol = getattr(kart, "raf_omru_tolerans", None) or "±%10"
    # Miktar Tayini hedefini bitmiş testlerden bul
    hedef = None; birim = ""
    for test in kart.testler:
        if "miktar tayini" in _norm_basit(test.ad) and test.spesifikasyon.hedef_deger:
            hedef = test.spesifikasyon.hedef_deger
            birim = test.spesifikasyon.birim or ""
            break
    if hedef is None:
        return
    yeni_spek = _miktar_spek_uret(hedef, tol, birim)
    for row in t.rows:
        if len(row.cells) < 2:
            continue
        sol = _norm_basit(row.cells[0].text)
        if "miktar tayini" in sol:
            _hucre_temizle(row.cells[-1])
            p = row.cells[-1].paragraphs[0]
            r = p.add_run(yeni_spek)
            r.font.name = "Times New Roman"
            break


def _doldur_tablo89_eski(doc, proje: ProjeVerisi) -> None:

    def _satirlari_uret(raf_omru: bool, tol: str):
        yildiz = "*" if raf_omru else ""
        satirlar = []  # (sol, sag)
        imp_eklendi = False
        enan_eklendi = False

        def _ilgili_bilesikler_ekle():
            """İlgili Bileşikler başlığı + alt satırlar (etkenden). Tek etkende
            'e Ait' ara başlığı KULLANILMAZ (girdi biçimiyle birebir)."""
            if not any(em.impuriteler for em in etkenler):
                return
            satirlar.append(("İlgili Bileşikler", ""))
            cok_etken = sum(1 for em in etkenler if em.impuriteler) > 1
            for em in etkenler:
                if not em.impuriteler:
                    continue
                if cok_etken:
                    satirlar.append((f"{em.ad}'e Ait", ""))
                for imp in em.impuriteler:
                    a = imp.ad if imp.ad.startswith(("—", "-", "–")) else f"— {imp.ad}"
                    satirlar.append((a, imp.limit_metni or ""))

        def _enantiomerik_ekle():
            """Enantiomerik İmpurite başlığı + alt satırlar (etkenden)."""
            if not any(getattr(em, "enantiomerik", None) for em in etkenler):
                return
            satirlar.append(("Enantiomerik İmpurite", ""))
            for em in etkenler:
                for imp in (getattr(em, "enantiomerik", None) or []):
                    a = imp.ad if imp.ad.startswith(("—", "-", "–")) else f"— {imp.ad}"
                    satirlar.append((a, imp.limit_metni or ""))

        for test in bitmis:
            ad = test.ad
            n = _norm_basit(ad)
            spek = test.spesifikasyon.spesifikasyon_metni or test.spesifikasyon.metni_olustur()
            # Mikrobiyolojik: başlık + alt satırlar. Ondan ÖNCE İlgili Bileşikler
            # ve Enantiomerik henüz eklenmediyse ekle (girdi sırası: ...Enan, Mikro).
            if test.mikrobiyolojik:
                if not imp_eklendi:
                    _ilgili_bilesikler_ekle(); imp_eklendi = True
                if not enan_eklendi:
                    _enantiomerik_ekle(); enan_eklendi = True
                satirlar.append(("Mikrobiyolojik Kontrol", ""))
                for et, sp in (test.alt_satirlar or []):
                    satirlar.append((et, sp))
                continue
            # Ağırlık Tekdüzeliği/Sapması: başlık (boş) + 2 alt satır
            if test.aciklama_etiketi:
                satirlar.append((ad, ""))
                satirlar.append((test.aciklama_etiketi, test.aciklama_spek))
                if test.aciklama2_etiketi:
                    satirlar.append((test.aciklama2_etiketi, test.aciklama2_spek))
                continue
            # Boyar Madde: başlık (boş) + alt satırlar (Titanyum dioksit vb.)
            if "boyar madde" in n and (test.alt_satirlar or not spek):
                satirlar.append((ad, ""))
                for et, sp in (test.alt_satirlar or []):
                    a = et if et.startswith(("—", "-", "–")) else f"-{et}"
                    satirlar.append((a, sp))
                continue
            # İlgili Bileşikler bitmiş listede AYRI test olarak varsa: burada ekle
            if "ilgili bilesik" in n:
                if not imp_eklendi:
                    _ilgili_bilesikler_ekle(); imp_eklendi = True
                continue
            if "enantiomerik" in n:
                if not enan_eklendi:
                    _enantiomerik_ekle(); enan_eklendi = True
                continue
            # Miktar Tayini: raf ömründe tolerans uygula
            if raf_omru and "miktar tayini" in n and test.spesifikasyon.hedef_deger:
                spek = _miktar_spek_uret(test.spesifikasyon.hedef_deger, tol, test.spesifikasyon.birim)
            satirlar.append((ad, spek))

        # Döngü mikrobiyolojik olmadan bittiyse İlgili/Enantiomerik'i sona ekle
        if not imp_eklendi:
            _ilgili_bilesikler_ekle()
        if not enan_eklendi:
            _enantiomerik_ekle()
        return satirlar

    raf_omru_ekle = getattr(kart, "tablo89_ekle", True)
    hedefler = [(8, False, kart.serbest_birakma_tolerans)]
    if raf_omru_ekle:
        hedefler.append((9, True, kart.raf_omru_tolerans))
    else:
        # Tablo 9 üretilmeyecekse şablondaki TASLAK satırları temizle
        # (boş alanda taslak metni bırakma kuralı).
        t9 = _tablo_basliga_gore(doc, 9)
        if t9 is not None:
            for ri in range(1, len(t9.rows)):
                for cell in t9.rows[ri].cells:
                    _hucre_temizle(cell)
    for tablo_no, raf, tol in hedefler:
        t = _tablo_basliga_gore(doc, tablo_no)
        if t is None:
            continue
        satirlar = _satirlari_uret(raf, tol)
        if not satirlar:
            continue
        idxler = _veri_satirlarini_ayarla(t, 1, len(satirlar))
        for ri, (sol, sag) in zip(idxler, satirlar):
            cells = t.rows[ri].cells
            hucre_yaz(cells[0], sol)
            hucre_yaz(cells[-1], sag)
        _tablo_genislik_duzelt(t)


def _norm_basit(s: str) -> str:
    tr = str.maketrans({"ı": "i", "İ": "i", "ş": "s", "ç": "c", "ö": "o",
                        "ü": "u", "ğ": "g", "I": "i"})
    return (s or "").translate(tr).lower()


def _kucuk(metin: str) -> str:
    """Türkçe-güvenli küçük harf ('İ'.lower() combining-dot sorununu önler)."""
    if not metin:
        return ""
    return (metin.replace("İ", "i").replace("I", "ı")
            .replace("Ş", "ş").replace("Ğ", "ğ")
            .replace("Ü", "ü").replace("Ö", "ö").replace("Ç", "ç")
            .lower())


def _dikey_ok_xml(yukseklik_emu=180000):
    """Aşağı bakan dikey ok. Bounding box dikey-ince olduğundan çizgi dikeydir."""
    import random
    sid = random.randint(1000, 9999999)
    w = 12700  # ok çizgisi kalınlığı = bounding box genişliği
    return (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:rPr/><w:drawing>'
        '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{w}" cy="{yukseklik_emu}"/>'
        f'<wp:docPr id="{sid}" name="dok{sid}"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:cNvCnPr/><wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="{yukseklik_emu}"/></a:xfrm>'
        '<a:prstGeom prst="straightConnector1"><a:avLst/></a:prstGeom>'
        '<a:ln w="12700" cap="flat">'
        '<a:solidFill><a:srgbClr val="000000"/></a:solidFill>'
        '<a:tailEnd type="triangle" w="med" len="med"/>'
        '</a:ln>'
        '</wps:spPr><wps:bodyPr/>'
        '</wps:wsp></a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'
    )


def _yatay_ok_xml(genislik_emu, yukseklik_emu):
    """Sağa bakan yatay ok. Bounding box yatay-ince olduğundan çizgi yataydır."""
    import random
    sid = random.randint(1000, 9999999)
    h = 12700  # ok çizgisi kalınlığı = bounding box yüksekliği
    return (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:rPr/><w:drawing>'
        '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{genislik_emu}" cy="{h}"/>'
        f'<wp:docPr id="{sid}" name="ok{sid}"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:cNvCnPr/><wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{genislik_emu}" cy="0"/></a:xfrm>'
        '<a:prstGeom prst="straightConnector1"><a:avLst/></a:prstGeom>'
        '<a:ln w="12700" cap="flat">'
        '<a:solidFill><a:srgbClr val="000000"/></a:solidFill>'
        '<a:tailEnd type="triangle" w="med" len="med"/>'
        '</a:ln>'
        '</wps:spPr><wps:bodyPr/>'
        '</wps:wsp></a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'
    )


def _kutu_sekli_xml(metin, genislik_emu, yukseklik_emu, dolgu="D9E2F3", oklu=False,
                    geom="roundRect"):
    """
    Bir kutu şekli (geom: roundRect/rect) + içinde ortalı metin üreten DrawingML
    XML'i döndürür. Metin Word formatında (w:txbxContent içinde w:p) yazılır.
    """
    import html, random
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    sat = metin.split("\n")
    paragraflar = ""
    for s in sat:
        paragraflar += (
            f'<w:p xmlns:w="{W}"><w:pPr><w:jc w:val="left"/>'
            '<w:spacing w:after="0" w:line="200" w:lineRule="exact"/></w:pPr>'
            '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            '<w:sz w:val="14"/><w:szCs w:val="14"/>'
            '<w:color w:val="000000"/></w:rPr>'
            f'<w:t xml:space="preserve">{html.escape(s)}</w:t></w:r></w:p>'
        )
    sid = random.randint(1000, 9999999)
    return (
        f'<w:r xmlns:w="{W}">'
        '<w:rPr/><w:drawing>'
        '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{genislik_emu}" cy="{yukseklik_emu}"/>'
        f'<wp:docPr id="{sid}" name="kutu{sid}"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:cNvSpPr/><wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{genislik_emu}" cy="{yukseklik_emu}"/></a:xfrm>'
        f'<a:prstGeom prst="{geom}"><a:avLst/></a:prstGeom>'
        '<a:noFill/>'
        '<a:ln w="6350"><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>'
        '</wps:spPr><wps:txbx><w:txbxContent>'
        + paragraflar +
        '</w:txbxContent></wps:txbx>'
        '<wps:bodyPr rot="0" anchor="ctr" anchorCtr="0" lIns="36000" tIns="9000" '
        'rIns="18000" bIns="9000"/>'
        '</wps:wsp></a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'
    )


def _doldur_akis_semasi(doc, proje: ProjeVerisi) -> None:
    """
    'Proses Akış Diyagramı' başlığı altına 4 sütunlu akış şeması çizer.
    Önce şablondaki eski (textbox'lı örnek) şemayı siler.
    Sütunlar: Hammaddeler | Operasyon kutuları (şekil + ↓ ok) | İPK | Kimyasal.
    """
    try:
        from core.akis_semasi import akis_semasi_hazirla
        veri = akis_semasi_hazirla(proje)
        kutular = veri.get("operasyonlar", [])
        hammadde_kutu = veri.get("hammaddeler", [])
    except Exception:
        kutular = []
        hammadde_kutu = []
    if not kutular:
        return
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement, parse_xml
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    paralar = list(doc.paragraphs)
    bas_idx = None
    for i, p in enumerate(paralar):
        tx = p.text.strip().lower()
        if "proses akış" in tx or "akış diyagram" in tx:
            bas_idx = i
            break
    if bas_idx is None:
        return

    capa = paralar[bas_idx]._p
    capa_parent = paralar[bas_idx]._parent

    # --- Başlık öncesine sayfa sonu ekle (akış şeması sayfa başından başlasın) ---
    from docx.enum.text import WD_BREAK
    onceki = capa.getprevious()
    # zaten sayfa sonu yoksa ekle
    pb = OxmlElement("w:p")
    r_pb = OxmlElement("w:r")
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    r_pb.append(br); pb.append(r_pb)
    capa.addprevious(pb)

    # --- Eski şema şekillerini sil (başlık paragrafı + sonraki paragraflardaki
    #     drawing/pict öğeleri; metin korunur) ---
    def _sekilleri_temizle(p_el):
        for tag in ("w:drawing", "w:pict"):
            for d in p_el.findall(".//" + qn(tag)):
                # drawing/pict bir w:r içinde; tüm run'ı kaldır
                r = d.getparent()
                while r is not None and not r.tag.endswith("}r"):
                    r = r.getparent()
                if r is not None and r.getparent() is not None:
                    r.getparent().remove(r)
                elif d.getparent() is not None:
                    d.getparent().remove(d)

    _sekilleri_temizle(capa)  # başlık paragrafındaki eski şekiller

    el = capa.getnext()
    sayac = 0
    while el is not None and sayac < 8:
        sayac += 1
        if el.tag.endswith("}tbl"):
            txbx = el.findall(".//" + qn("w:txbxContent"))
            drawing = el.findall(".//" + qn("w:drawing"))
            pict = el.findall(".//" + qn("w:pict"))
            if len(txbx) > 3 or len(drawing) > 3 or len(pict) > 3:
                sil = el; el = el.getnext()
                sil.getparent().remove(sil)
                continue
            else:
                break
        elif el.tag.endswith("}p"):
            ptxt = "".join(el.itertext()).strip().lower()
            draw = el.findall(".//" + qn("w:drawing"))
            pict = el.findall(".//" + qn("w:pict"))
            if draw or pict:
                _sekilleri_temizle(el)  # paragraftaki eski şekilleri sil
            if ptxt and ("tablo" in ptxt or "kapsanan" in ptxt or "risk" in ptxt
                         or "stabilite" in ptxt):
                break
        el = el.getnext()

    def _hucre_yaz(cell, satirlar, *, bold=False, ortala=True):
        # Kullanıcı Word'de elle ayarlamak zorunda kalmasın: hücre metni ÜSTE
        # hizalı (top) ve metin KAYDIRMALI (word-wrap açık, noWrap kaldırılır).
        tcPr = cell._tc.get_or_add_tcPr()
        va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "top")
        tcPr.append(va)
        # noWrap varsa kaldır (metin kaydırma açık kalsın)
        nw = tcPr.find(qn("w:noWrap"))
        if nw is not None:
            tcPr.remove(nw)
        ilk = True
        for metin in satirlar:
            p = cell.paragraphs[0] if ilk else cell.add_paragraph()
            ilk = False
            if ortala:
                p.alignment = 1
            run = p.add_run(metin)
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _yuk_hesapla(metin):
        KAR_SIGAR = 30
        gorsel_satir = 0
        for satir in metin.split("\n"):
            gorsel_satir += max(1, -(-len(satir) // KAR_SIGAR))
        return 230000 + max(1, gorsel_satir) * 170000

    def _kutu_paragraf(cell, metin, ilk, oklu=False, hizala_yukseklik=None, geom="roundRect"):
        """Hücreye bir şekil-kutu ekler (ok YOK). geom: roundRect/rect."""
        p = cell.add_paragraph() if (cell.paragraphs[0].runs or not ilk) else cell.paragraphs[0]
        p.alignment = 1
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        yuk = _yuk_hesapla(hizala_yukseklik if hizala_yukseklik else metin)
        xml = _kutu_sekli_xml(metin, 1330000, yuk, dolgu=None, geom=geom)
        p._p.append(parse_xml(xml))

    def _bos_hiza(cell, hizala_metin, ilk):
        """Sütun 1'de hizalama için, verilen metin yüksekliğinde boş paragraf."""
        p = cell.add_paragraph() if (cell.paragraphs[0].runs or not ilk) else cell.paragraphs[0]
        p.alignment = 1
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        # görünmez ama yer kaplayan boşluk: yükseklik kadar satır
        yuk = _yuk_hesapla(hizala_metin)
        # boş kutu yerine sadece boşluk paragrafı (yükseklik için satır)
        satir_say = max(1, round(yuk / 170000))
        r = p.add_run("\n" * (satir_say - 1) if satir_say > 1 else " ")
        r.font.size = Pt(7)

    def _bos_satir(cell):
        """Hücreyi boş bırak (tek küçük boşluk paragrafı)."""
        p = cell.paragraphs[0]
        p.alignment = 1
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(" ")
        r.font.size = Pt(6)

    def _ic_kenarlik_gizle(tablo):
        """Tablonun iç ve içerik satır kenarlıklarını gizler; sadece dış çerçeve
        ve başlık altı çizgisi görünür (satırlar tek kutu gibi görünür)."""
        rows = tablo.rows
        for ri, row in enumerate(rows):
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                # mevcut tcBorders varsa temizle
                for eski in tcPr.findall(qn("w:tcBorders")):
                    tcPr.remove(eski)
                tcb = OxmlElement("w:tcBorders")
                # başlık satırı (ri=0): alt çizgi görünür; diğerleri gizli iç çizgi
                for kenar in ("top", "bottom", "left", "right"):
                    b = OxmlElement(f"w:{kenar}")
                    gorunur = False
                    if kenar == "left" or kenar == "right":
                        gorunur = True  # dış yan çerçeve
                    if ri == 0 and kenar == "bottom":
                        gorunur = True  # başlık altı çizgi
                    if ri == 0 and kenar == "top":
                        gorunur = True
                    if ri == len(rows) - 1 and kenar == "bottom":
                        gorunur = True  # en alt çerçeve
                    b.set(qn("w:val"), "single" if gorunur else "nil")
                    if gorunur:
                        b.set(qn("w:sz"), "4"); b.set(qn("w:color"), "000000")
                    tcb.append(b)
                tcPr.append(tcb)

    # --- Tablo iskeleti (kenarlıksız, tek satır) ---
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    st = OxmlElement("w:tblStyle"); st.set(qn("w:val"), "TableGrid"); tblPr.append(st)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); tblPr.append(jc)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), "0"); tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    borders = OxmlElement("w:tblBorders")
    for kenar in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{kenar}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    tbl.append(tblPr)
    grid = OxmlElement("w:tblGrid")
    for w in (2300, 2300, 2300, 2300):
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    tbl.append(grid)
    capa.addnext(tbl)
    tablo = Table(tbl, capa_parent)

    # başlık satırı
    bsr = tablo.add_row().cells
    _hucre_yaz(bsr[0], ["Hammaddeler"], bold=True)
    _hucre_yaz(bsr[1], ["Üretim Aşaması"], bold=True)
    _hucre_yaz(bsr[2], ["İPK Testleri"], bold=True)
    _hucre_yaz(bsr[3], ["Kimyasal Analizler"], bold=True)

    # --- TEK içerik satırı: 4 hücre, her hücrede dikey kutular (ok YOK) ---
    cells = tablo.add_row().cells
    for ust_cell in cells:
        va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "top")
        ust_cell._tc.get_or_add_tcPr().append(va)

    # Sütun 0: Hammadde kutuları — DİKDÖRTGEN (köşeli) şekil
    for j, ham in enumerate(hammadde_kutu):
        metin = ham["metin"] if isinstance(ham, dict) else ham
        _kutu_paragraf(cells[0], metin, ilk=(j == 0), geom="rect")

    # Sütun 1: Üretim Aşaması — BOŞ (kullanıcı kendisi dolduracak)

    # Sütun 2: İPK ve Sütun 3: Kimyasal — yuvarlatılmış kutu (değiştirilmedi)
    for i, k in enumerate(kutular):
        ipk_metin = "\n".join("- " + t for t in k["ipk_testleri"]) if k["ipk_testleri"] else "—"
        _kutu_paragraf(cells[2], ipk_metin, ilk=(i == 0), geom="roundRect")
        kim_metin = "\n".join("- " + t for t in k["kimyasal_testler"]) if k["kimyasal_testler"] else "—"
        _kutu_paragraf(cells[3], kim_metin, ilk=(i == 0), geom="roundRect")
    return


def _doldur_akis_semasi_ESKI(doc, proje: ProjeVerisi) -> None:
    return
    # (eski kod aşağıda kullanılmıyor)
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    kutular = []
    paralar = list(doc.paragraphs)
    bas_idx = None
    capa = paralar[0]._p
    capa_parent = paralar[0]._parent

    def _hucre_yaz(cell, satirlar, *, bold=False, ortala=True, kutu=False):
        """Bir tablo hücresine metin(ler) yazar; kutu=True ise gölgeli/kenarlıklı."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        if kutu:
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "D9E2F3"); tcPr.append(shd)  # açık mavi kutu
        va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "center"); tcPr.append(va)
        ilk = True
        for metin in satirlar:
            if ilk:
                p = cell.paragraphs[0]; ilk = False
            else:
                p = cell.add_paragraph()
            if ortala:
                p.alignment = 1
            run = p.add_run(metin)
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(10 if not kutu else 11)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Tablo: her aşama için 1 satır (operasyon kutusu) + ok satırı
    n_satir = len(kutular)
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    st = OxmlElement("w:tblStyle"); st.set(qn("w:val"), "TableGrid"); tblPr.append(st)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); tblPr.append(jc)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), "0"); tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    # sadece iç kenarlık yok; kutuları hücre gölgesiyle göstereceğiz
    borders = OxmlElement("w:tblBorders")
    for kenar in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{kenar}")
        b.set(qn("w:val"), "none"); borders.append(b)
    tblPr.append(borders)
    tbl.append(tblPr)
    grid = OxmlElement("w:tblGrid")
    for w in (2600, 2200, 2300, 2300):  # Hammadde | Operasyon | IPK | Kimyasal
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    tbl.append(grid)

    from docx.table import Table
    capa.addnext(tbl)
    tablo = Table(tbl, capa_parent)
    # başlık satırı
    bsr = tablo.add_row().cells
    _hucre_yaz(bsr[0], ["Hammaddeler"], bold=True)
    _hucre_yaz(bsr[1], ["Üretim Aşaması"], bold=True)
    _hucre_yaz(bsr[2], ["İPK Testleri"], bold=True)
    _hucre_yaz(bsr[3], ["Kimyasal Analizler"], bold=True)

    for ki, k in enumerate(kutular):
        cells = tablo.add_row().cells
        # sol: hammaddeler
        ham = ["- " + h for h in k["hammaddeler"]] if k["hammaddeler"] else [""]
        _hucre_yaz(cells[0], ham, ortala=False)
        # orta: operasyon kutusu (gölgeli) + alt ok
        ok = " ↓" if ki < len(kutular) - 1 else ""
        _hucre_yaz(cells[1], [k["operasyon"] + ok], bold=True, kutu=True)
        # sağ: IPK testleri (sadece ipk işaretliyse)
        ipk = ["- " + t for t in k["ipk_testleri"]] if k["ipk_testleri"] else [""]
        _hucre_yaz(cells[2], ipk, ortala=False)
        # sağ: kimyasal
        kim = ["- " + t for t in k["kimyasal_testler"]] if k["kimyasal_testler"] else [""]
        _hucre_yaz(cells[3], kim, ortala=False)


def _doldur_uretim_yontemi(doc, proje: ProjeVerisi) -> None:
    """
    Şablonda 'üretim prosesi ... açıklanmaktadır' ile 'Proses Akış Diyagramı'
    arasındaki örnek paragrafları siler, kullanıcının üretim adımlarıyla
    değiştirir. Her adım: başlık (kalın) + açıklama + (varsa) parametre tablosu.
    """
    adimlar = getattr(proje, "uretim_adimlari", None)
    if not adimlar:
        return
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    paralar = list(doc.paragraphs)
    bas_idx = son_idx = None
    for i, p in enumerate(paralar):
        t = p.text.strip().lower()
        if bas_idx is None and "üretim prosesi" in t and "açıklan" in t:
            bas_idx = i
        elif bas_idx is not None and ("proses akış" in t or "proses akis" in t):
            son_idx = i
            break
    if bas_idx is None or son_idx is None:
        return

    for p in paralar[bas_idx + 1:son_idx]:
        p._p.getparent().remove(p._p)

    # Aradaki şablon ÖRNEK tablolarını da sil (paragraflar arasındaki w:tbl'ler)
    bas_p_el = paralar[bas_idx]._p
    son_p_el = paralar[son_idx]._p
    el = bas_p_el.getnext()
    while el is not None and el is not son_p_el:
        sonraki = el.getnext()
        if el.tag.endswith("}tbl"):
            el.getparent().remove(el)
        el = sonraki

    capa_parent = paralar[bas_idx]._parent
    son_el = paralar[bas_idx]._p

    def _para_ekle(sonra_el, metin, bold):
        yp = OxmlElement("w:p")
        sonra_el.addnext(yp)
        para = Paragraph(yp, capa_parent)
        # Başlık (bold) paragrafı: sonraki paragrafla (metniyle) AYNI sayfada kalsın
        if bold:
            pPr = yp.get_or_add_pPr()
            kn = OxmlElement("w:keepNext"); pPr.append(kn)
            kl = OxmlElement("w:keepLines"); pPr.append(kl)
        run = para.add_run(metin)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return yp

    def _tablo_ekle(sonra_el, satirlar):
        """2 sütunlu dar, ortalı parametre tablosu (taslaktaki gibi)."""
        tbl = OxmlElement("w:tbl")
        tblPr = OxmlElement("w:tblPr")
        st = OxmlElement("w:tblStyle"); st.set(qn("w:val"), "TableGrid"); tblPr.append(st)
        # ortalı
        jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); tblPr.append(jc)
        tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), "0"); tblW.set(qn("w:type"), "auto")
        tblPr.append(tblW)
        borders = OxmlElement("w:tblBorders")
        for kenar in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{kenar}")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
            borders.append(b)
        tblPr.append(borders)
        tbl.append(tblPr)
        # dar sütunlar (taslaktaki gibi: ~2223 / 2410)
        grid = OxmlElement("w:tblGrid")
        for w in (2223, 2410):
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
        tbl.append(grid)
        for sol, sag in satirlar:
            tr = OxmlElement("w:tr")
            for metin, gen in ((sol, 2223), (sag, 2410)):
                tc = OxmlElement("w:tc")
                tcPr = OxmlElement("w:tcPr")
                tcW = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(gen)); tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW); tc.append(tcPr)
                pp = OxmlElement("w:p")
                run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), "Times New Roman")
                rf.set(qn("w:hAnsi"), "Times New Roman"); rPr.append(rf)
                sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)
                col = OxmlElement("w:color"); col.set(qn("w:val"), "000000"); rPr.append(col)
                run.append(rPr)
                wt = OxmlElement("w:t"); wt.text = metin; run.append(wt)
                pp.append(run); tc.append(pp); tr.append(tc)
            tbl.append(tr)
        sonra_el.addnext(tbl)
        return tbl

    son_el = _para_ekle(son_el, "", False)
    for adim in adimlar:
        # adım (baslik, aciklama) veya (baslik, aciklama, tablo) olabilir
        baslik = adim[0]
        aciklama = adim[1] if len(adim) > 1 else ""
        tablo = adim[2] if len(adim) > 2 else []
        son_el = _para_ekle(son_el, baslik, True)
        son_el = _para_ekle(son_el, aciklama, False)
        if tablo:
            son_el = _tablo_ekle(son_el, tablo)
        son_el = _para_ekle(son_el, "", False)


def _doldur_uretim_yontemi_eski(doc, proje):
    pass


def _ortak_doldur(doc, proje: ProjeVerisi, rapor: bool) -> None:
    proje._formul_notlari_yazildi = False
    belgede_degistir(doc, _placeholder_eslemeleri(proje, rapor))
    _revizyon_no_guncelle(doc, proje.dokuman.revizyon_no)
    _revizyon_tarihi_guncelle(doc, proje.dokuman.revizyon_tarihi)
    _doldur_formul(doc, proje)
    _doldur_kapsanan(doc, proje)
    _doldur_risk(doc, proje)
    _doldur_proses_param(doc, proje)
    _doldur_ekipman(doc, proje)
    _doldur_uretim_yontemi(doc, proje)
    _doldur_akis_semasi(doc, proje)
    _doldur_spek(doc, proje)
    _doldur_ipk(doc, proje)
    _doldur_tablo89(doc, proje)
    _doldur_numune(doc, proje)
    belgede_degistir(doc, _placeholder_eslemeleri(proje, rapor))
    _doldur_formul(doc, proje)
    _doldur_kapsanan(doc, proje)
    _doldur_risk(doc, proje)
    _doldur_proses_param(doc, proje)
    _doldur_ekipman(doc, proje)
    _doldur_uretim_yontemi(doc, proje)
    _doldur_akis_semasi(doc, proje)
    _doldur_spek(doc, proje)
    _doldur_ipk(doc, proje)
    _doldur_tablo89(doc, proje)
    _doldur_numune(doc, proje)


def _turetilmis_testlerle(proje: ProjeVerisi):
    """
    Context manager: otomatik_turet açıksa, spek_karti.testler (bitmiş ürün
    listesi) kural motoruyla tüm aşamalara dağıtılır. Çıkışta orijinal liste
    geri yüklenir (kullanıcının girdiği liste bozulmaz).
    """
    import contextlib

    @contextlib.contextmanager
    def _ctx():
        kart = proje.spek_karti
        if not getattr(kart, "otomatik_turet", False):
            yield
            return
        from core.kural_motoru import turet
        orijinal = kart.testler
        kart._bitmis_urun_testleri = orijinal  # Tablo 8/9 bunu kullanır
        ops = proje.urun_formu.operasyonlar
        kart.testler = turet(orijinal, kart.etkin_maddeler, ops,
                             cift_katman=getattr(kart, "cift_katman", False),
                             tablet_ipk=getattr(kart, "tablet_ipk", {}),
                             ozel_test_kurallari=getattr(kart, "ozel_test_kurallari", {}))
        try:
            yield
        finally:
            kart.testler = orijinal
            kart._bitmis_urun_testleri = None
    return _ctx()


def pvp_uret(proje: ProjeVerisi, cikti_yolu: str | Path) -> Path:
    with _turetilmis_testlerle(proje):
        doc = Document(str(_sablon_yolu("PVP_sablon.docx")))
        _ortak_doldur(doc, proje, rapor=False)
        _tum_belge_fontu(doc)
        yol = Path(cikti_yolu)
        doc.save(str(yol))
    return yol


def pvr_uret(proje: ProjeVerisi, cikti_yolu: str | Path, veri_uret: bool = True,
             tohum: int | None = None) -> Path:
    with _turetilmis_testlerle(proje):
        if veri_uret:
            vu.tum_testleri_uret(proje.spek_karti.testler, tohum=tohum)
        # PVR, TEMIZ PVP şablonundan üretilir; sonuç tabloları (Bölüm 11) temiz eklenir.
        doc = Document(str(_sablon_yolu("PVP_sablon.docx")))
        _ortak_doldur(doc, proje, rapor=True)
        _doldur_sonuclar(doc, proje)
        _doldur_genel_degerlendirme(doc, proje)
        _tum_belge_fontu(doc)
        yol = Path(cikti_yolu)
        doc.save(str(yol))
    return yol
